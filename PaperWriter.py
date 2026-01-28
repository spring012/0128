from Interface import PaperWriterFramework, Locker, PaperInfo, PaperInfoAll, ChapterItem, BaseWriter
from LargeModel import OnlineModel
from DataProcess import ReferenceDataBase
from Retriever import Retriever
import Prompt
import Global

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List
import re
import random
import torch
# import threading
import json
import win32com.client as win32
import os
import datetime
import requests
import time
from langdetect import detect
from langchain.prompts import ChatPromptTemplate
from langchain.schema.output_parser import StrOutputParser
import matplotlib.pyplot as plt
import numpy as np
from docx.shared import Inches
from plantuml import PlantUML  # 确保导入 PlantUML 库

def clean_content(content: str, is_English=False):
    if not is_English:
        content = re.sub(' ', '', content)
    content = re.sub('#', '', content)
    content = re.sub('[*]', '', content)
    content = re.sub('-', '', content)
    # content = re.sub('\n', '', content)
    return content

def seq_segment(content: str):
    list_content = []
    for item in content.split('\n'):
        if len(item) != 0:
            list_content.append(item)
    return list_content

def check_content(content: str, partten: list):
    for item in partten:
        if item in content:
            return True
    return False

class AbstractWriter(BaseWriter):
    def __init__(self, llm: OnlineModel, ref_db: ReferenceDataBase, paper_info: PaperInfo, data_volume: int, log_path='') -> None:
        """
        字数控制：
        5000 1级
        10000 2级
        20000 23级
        30000 23级
        useThreeLevel用于强制控制标题等级
        """
        super().__init__(llm, paper_info, log_path)
        self.UseTreeLevel = paper_info.useThreeLevel or self.word_count >= 20000
        self.customOutline = paper_info.customOutline
        
        self.retriever = ref_db.embedding_database.as_retriever(search_kwargs={"k": data_volume}, search_type="mmr")
        self.number_to_chinese = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九', 10:'十'}
        self.language = paper_info.language
        self.is_English=paper_info.isEnglish
        # print(self.language,self.is_English)
        self.log = []
        

    # def __del__(self):
    #     if self.log_path != '':
    #         with open(self.log_path + self.title + '_摘要.json', 'w', encoding='utf-8') as f:
    #             json.dump({'OUTPUT': self.log}, f, ensure_ascii=False, indent=4)

    def LegalCheck(self, min_abstract_word:int) -> bool:
        try:
            prompt = ChatPromptTemplate.from_template(Prompt.abstract_template.prompt)
            context = []

            # 如果有非法信息直接抛出异常
            legal_chain = prompt | OnlineModel(llm_name='Zhipu-Chinese').llm | StrOutputParser()
            legal_check = legal_chain.invoke({'context': context, 
                                    'form': Prompt.abstract_template.form,
                                    'category':self.category, 
                                    'min_abstract_word': min_abstract_word + 100,
                                    'title':self.title,
                                    'customOutline': "" if self.customOutline == "" else f"\nPaper Outline：\n{self.customOutline}",
                                    'language': self.language}
                                    )
            return True
        except Exception as e:
            return False
        
    def GetAbstract(self, min_abstract_word:int) -> tuple:
        prompt = ChatPromptTemplate.from_template(Prompt.abstract_template.prompt)
        chain = prompt | self.llm_abstract | StrOutputParser()
        context = []

        # print('开始做RAG')
        # context = self.retriever.invoke(self.category + self.title)
        # print('RAG结束')
        # context = self.limit_input_length(context)

        # context.append(f"Paper Outline：\n{self.customOutline}")

        abstract = chain.invoke({'context': context, 
                                'form': Prompt.abstract_template.form,
                                'category':self.category, 
                                'min_abstract_word': min_abstract_word + 100,
                                'title':self.title,
                                'customOutline': "" if self.customOutline == "" else f"\nPaper Outline：\n{self.customOutline}",
                                'language': self.language}
                                )
        print('摘要生成完成')
        if self.log_path != '':
            self.__write_log(content=abstract, documents=context)
        
        try:
            match = re.search(r"(\{.*\})", abstract, re.DOTALL)
            if match:
                content = match.groups()[0]
            cc = content.strip()
            if cc.startswith('json```'):
                cc = cc.replace('json```', '', 1)
            if cc.endswith('```'):
                cc = cc[:-3]
            struct_data = json.loads(cc.strip())
            self.abstract = (self.reduce_AIGC(struct_data.get('Abstract')), struct_data.get('Keywords'))
            print('摘要解析完成')
            return self.abstract
        except Exception as e:
            print(f"The response is not a valid JSON format: {e}")
            return (None, None)
    
    def GetOutline(self) -> List[ChapterItem]:
        if not hasattr(self, 'abstract'):
            return None

        prompt = ChatPromptTemplate.from_template(Prompt.outlines_template.two_level_outlines)
        
        chain = prompt | self.llm | StrOutputParser()
        outline = chain.invoke({'abstract':self.abstract[0], 
                        'customOutline': "" if self.customOutline == "" else f"Refer to this outline:\n{self.customOutline}",
                        'title':self.title, 
                        'form': Prompt.outlines_template.two_level_form,
                        'language': self.language
                        })
        print('提纲生成完成')
        
        if self.log_path != '':
            self.__write_log(content=outline)
        
        try:
            match = re.search(r"(\{.*\})", outline, re.DOTALL)
            if match:
                content = match.groups()[0]
            cc = content.strip()
            if cc.startswith('json```'):
                cc = cc.replace('json```', '', 1)
            if cc.endswith('```'):
                cc = cc[:-3]
            two_level_headings = json.loads(cc.strip())['Outline']
        except Exception as e:
            print(f"The response is not a valid JSON format: {e}")
            return None
        
        for index, (chapter, sub_titles) in enumerate(two_level_headings.items()):
            if index == len(two_level_headings) - 1:
                del two_level_headings[chapter]
                if self.is_English:
                    two_level_headings['Chapter ' + str(index + 1) + ' Conclusion'] = []
                else:
                    two_level_headings['第' + self.number_to_chinese[index + 1] + '章 总结'] = []
                break
        
        if self.UseTreeLevel:
            three_level_heandings = []
            info_format = []
            repeat_content = set()
            for index, (chapter, sub_titles) in enumerate(two_level_headings.items()):
                if index == len(two_level_headings) - 1:
                    break
                for item in sub_titles:
                    repeat_content.add(item)
                    info_format.append({
                        'abstract': self.abstract[0],
                        'customOutline': "" if self.customOutline == "" else f"Refer to this outline:\n{self.customOutline}",
                        'title': self.title, 'chapter': chapter, 'part': item, 'language': self.language, 'form': Prompt.outlines_template.three_level_form,})
            
            prompt = ChatPromptTemplate.from_template(Prompt.outlines_template.three_level_outlines)
            chain = prompt | self.llm | StrOutputParser()
            three_level_outline = chain.batch(info_format)
            print('三级提纲生成完成')
            for item in three_level_outline:
                try:
                    match = re.search(r"(\{.*\})", item, re.DOTALL)
                    if match:
                        content = match.groups()[0]
                    cc = content.strip()
                    if cc.startswith('json```'):
                        cc = cc.replace('json```', '', 1)
                    if cc.endswith('```'):
                        cc = cc[:-3]
                    three_level_heanding = json.loads(cc)
                    three_level_heandings.append(three_level_heanding['Subheadings'])
                except:
                    three_level_heandings.append([])
                    
        outlines = {}
        info_format = []
        sec_index = 0
        for chapter, sub_titles in two_level_headings.items():
            outlines[chapter] = {}
            info_format.append({'abstract':self.abstract[0],
                                'customOutline': "" if self.customOutline == "" else f"Refer to this outline:\n{self.customOutline}",
                                'title':self.title, 'part':chapter.split(' ')[-1], 'language': self.language})

            if self.word_count <= 5000:
                continue

            for sub_title in sub_titles:
                outlines[chapter][sub_title] = []
                info_format.append({'abstract':self.abstract[0],
                                    'customOutline': "" if self.customOutline == "" else f"Refer to this outline:\n{self.customOutline}",
                                    'title':self.title, 'part':sub_title, 'language': self.language})

                if self.UseTreeLevel:
                    outlines[chapter][sub_title] = []
                    for subsub_title in three_level_heandings[sec_index]:
                        if subsub_title not in repeat_content:
                            outlines[chapter][sub_title].append(subsub_title)
                            info_format.append({'abstract':self.abstract[0],
                                                'customOutline': "" if self.customOutline == "" else f"Refer to this outline:\n{self.customOutline}",
                                                'title':self.title, 'part':subsub_title, 'language': self.language})
                    if len(outlines[chapter][sub_title]) == 1:
                        outlines[chapter][sub_title] = []
                    sec_index += 1
 
        prompt = ChatPromptTemplate.from_template(Prompt.outlines_template.generate_info)
        chain = prompt | self.llm | StrOutputParser()

        info = chain.batch(info_format)
        print('小标题解释生成完成')

        if self.log_path != '':
            for item in info:
                self.__write_log(item)

        outlines_format = []
        info_index = 0
        for index, (part, subparts) in enumerate(outlines.items()):
            first_level = ChapterItem(title=part, info=info[info_index], sub=[], referenceFileList=[])
            info_index += 1

            for subpart, subsubparts in subparts.items():
                second_level = ChapterItem(title=subpart, info=info[info_index], sub=[], referenceFileList=[])
                info_index += 1

                for subsubpart in subsubparts:
                    second_level.sub.append(ChapterItem(title=subsubpart, info=info[info_index], sub=[], referenceFileList=[]))
                    info_index += 1

                first_level.sub.append(second_level)
            outlines_format.append(first_level)
        print('提纲解析完成')

        return outlines_format

    def __write_log(self, content: str, documents=[]):
        context = []
        if len(documents) != 0:
            for item in documents:
                context.append((item.page_content, item.metadata))

        self.log.append({'context': context, 'content': content})

class PaperWriter(BaseWriter):
    def __init__(self, llm: OnlineModel, retriever: Retriever, paper_info: PaperInfoAll, log_path='') -> None:
        """
        字数控制：
        5000    1级，   max_table 2,                                原生分段（0.8删顺序词）
        10000   2级，   max_table 3, 默认加目录，       章末分页，  段落合并（0.5删顺序词）
        20000   23级，  max_table 5, 默认加目录和英摘， 章末分页，  段落合并（0.5删顺序词）
        30000   23级，  max_table 7, 默认加目录和英摘， 章末分页，  段落合并（0.5删顺序词），  多段落
        """
        super().__init__(llm, paper_info, log_path)
        self.ref_databse = retriever.url_proprietary_sentence_databse
        self.local_retriever = retriever.local_chapter_retriever
        self.user_retrievers = retriever.user_chapter_retrievers
        self.full_retriever = retriever.full_chapter_retriever

        if self.word_count <= 5000:
            self.max_table_num = 2
        elif self.word_count <= 10000:
            self.max_table_num = 3
        elif self.word_count <= 20000:
            self.max_table_num = 5
        else:
            self.max_table_num = 7 

        # 没有3级标题，2级标题下的段落数
        self.single_subsubpart_min_passage = 1 if self.word_count < 30000 else 2
        self.single_subsubpart_max_passage = 1 if self.word_count < 30000 else 3
        # 有多个同级3级标题，其下的段落数
        self.multi_subsubpart_min_passage = 1
        self.multi_subsubpart_max_passage = 1 if self.word_count < 30000 else 2
        # 没有2，3级标题只有1级标题下的段落数
        self.single_chapter_min_passage = 1 if self.word_count < 30000 else 3
        self.single_chapter_max_passage = 1 if self.word_count < 30000 else 4

        self.number_to_chinese = {1:'一', 2:'二', 3:'三', 4:'四', 5:'五', 6:'六', 7:'七', 8:'八', 9:'九', 10:'十'}
        self.table_type = ['simulated data', 'theoretical', 'experiment steps']  # 可生成的数据表类型

        self.not_ref_text = ['本文', '本研究', '部分', '最后', '综上所述']  # 如果句子中有这些关键词，不进行引用
        self.sum_word = ['综上所述，', '所以，', '总的来说，', '根据以上分析，', '由此得出，', '综合上述论证，', '由此我们可以发现', '总而言之，', ''] # 替换综上所述的词

        self.ref_threshold = 0.5 # 超过这个相似度，文本就会被引用
        self.ref_probability = 0.3 # 低于ref_threshold情况下被引用的概率
        self.ref_continual_prob = 0.3 # 同一段中引用同一个文本的概率
        self.ref_second_chapter_prob = 0.5 # 第二章引用概率
        self.ref_subsequent_chapter_prob = 0.01 # 除了第一、二章外，其他章节的添加引用的概率
        self.max_ref_num = 3 # 一个括号里最大因用数
        self.ref_retrieve_num = 10 # 引用匹配时每次检索的文档数
        self.ref_papers = {'0': 0}

        self.delete_seq_prob = 0.8 if self.word_count <= 5000 else 0.5 # 去掉顺序词的概率

        self.EnglishAbstract = None
        self.thanks = None

        self.content_locker = Locker()
        self.log_locker = Locker()
        self.ref_locker = Locker()

        self.select_chapter_with_table= self.__select_chapter()

        self.word_distribution, self.words_per_section  = self.calculate_section_words() # 计算每个标题的目标字数
        # print(self.word_distribution)

    
    # def __del__(self):
        # if self.log_path != '':
        #     with open(self.log_path + self.title + '.json', 'w', encoding='utf-8') as f:
        #         json.dump({'OUTPUT': self.log_locker.content}, f, ensure_ascii=False, indent=4)

    def calculate_section_words(self) -> dict:
        """
        计算每个标题的目标字数，确保每个二级标题字数相等
        """
        word_distribution = {}
        
        # 计算基础每节字数（包括正文段落、表格描述和致谢）
        total_sections = (sum(1 if len(chapter.sub) == 0 else len(chapter.sub) for chapter in self.outlines) + 
                         len(self.select_chapter_with_table) + # 表格描述数
                         2) # 致谢和结论
        words_per_section = int((self.word_count - 500) / total_sections)  # -500是减去摘要的字数
        
        # 遍历每一章
        for part in self.outlines:
            for subpart in part.sub:
                if len(subpart.sub) == 0:
                    # 如果二级标题没有子标题，直接分配字数
                    word_distribution[subpart.title] = words_per_section
                else:
                    # 如果有三级标题，平分二级标题的字数
                    words_per_subsection = int(words_per_section / len(subpart.sub))
                    for subsub_part in subpart.sub:
                        word_distribution[subsub_part.title] = words_per_subsection

        return word_distribution, words_per_section

    def __select_chapter(self):
        select_chapter = []
        for index, part in enumerate(self.outlines):
            if index in (0, 1, len(self.outlines) - 1):
                continue
            if len(part.sub) == 0:
                select_chapter.append(part.title)

            for subpart in part.sub:
                if len(subpart.sub) == 0:
                    select_chapter.append(subpart.title)

                for subsubpart in subpart.sub:
                    select_chapter.append(subsubpart.title)
        
        if len(select_chapter) > self.max_table_num:
            select_chapter = random.sample(select_chapter, self.max_table_num)
        return select_chapter
        
    def GetFullPaper(self, forced_add_EnglishAbstract = False) -> dict:
        self.paper_content = {}
        threads = []

        if (forced_add_EnglishAbstract or self.word_count > 10000) and not self.is_English:
            """one_thread = threading.Thread(target=self.__add_EnglishAbstract, name='EnglishAbstract')
            threads.append(one_thread)
            one_thread.start()"""
            self.__add_EnglishAbstract()

        for index, part in enumerate(self.outlines):
            """one_thread = threading.Thread(target=self.__add_chapter, name=part.title, 
                                          args=(part, index + 1))
            threads.append(one_thread)
            one_thread.start()"""
            self.__add_chapter(part, index + 1)
            # if index == 0:
            #     break
        
        """one_thread = threading.Thread(target=self.__add_thanks, name='thanks')
        threads.append(one_thread)
        one_thread.start()

        for one_thread in threads:
            one_thread.join()"""
        self.__add_thanks()
        
        return self.paper_content
    
    def FullReduceAIGC(self):
        for index, (key, value) in enumerate(paper[0].items()):
            if index == '3':
                break
            for subkey, subvalue in value.items():
                if subkey == 'only_one_part':
                    for item in subsubvalue['text']:
                        # item = llm.invoke(prompt.format(context=item)).content
                        item = aigc(item)
                        print(item)
                        file.write(item)
                    continue
                for subsubkey, subsubvalue in subvalue.items():
                    for item in subsubvalue['text']:
                        # item = llm.invoke(prompt.format(context=item)).content
                        item = aigc(item)
                        print(item)
                        file.write(item)

    def __add_EnglishAbstract(self) -> tuple:
        prompt = ChatPromptTemplate.from_template(Prompt.translation_template.ChinsesToEnglish)
        chain = prompt | self.llm | StrOutputParser()

        abstract, key_word = chain.batch([self.abstract[0], self.abstract[1]])
        print('英文摘要生成完成')
        if self.log_path != '':
            self.write_log(content=abstract + '######' + key_word)

        key_word = key_word.split('\n')[0].split(':')[-1].strip()

        abstract = re.sub('\n', '', clean_content(abstract, is_English=True))
        key_word = clean_content(key_word, is_English=True)

        self.EnglishAbstract = (abstract, key_word)
        print('英文摘要解析完成')
        return self.EnglishAbstract
    
    def __add_passages(self, sub_part: ChapterItem, passage_num: int, chapter_index: int, repeat_content=''):
        # 获取当前标题的目标字数
        # print(sub_part)
        target_words = self.word_distribution.get(sub_part.title)
        if target_words is None:
            # print(f"Warning: 未找到 目标字数")
            target_words = self.words_per_section  # 设置默认值

        words_per_passage = int(target_words / passage_num)

        part_title = sub_part.title.split(' ')[-1]
        info = sub_part.info

        def add_single_passages():
            # 添加表格
            table = None
            image = None
            if sub_part.title in self.select_chapter_with_table:
                table = self.__add_table(topic=info, table_type=self.table_type[0], is_analysis=True)
                image = self.__add_image(table=table)
            # 添加正文
            if repeat_content == '':
                template = Prompt.single_passage_template.without_repeat
            else:
                template = Prompt.single_passage_template.with_repeat

            prompt = ChatPromptTemplate.from_template(template)
            generate_chain = prompt | self.llm | StrOutputParser()
            backup_generate_chain = prompt | OnlineModel(llm_name='GPT3.5').llm | StrOutputParser()

            if sub_part.title in self.user_retrievers:
                context = self.local_retriever.invoke(self.title + part_title + info) + self.user_retrievers[sub_part.title].invoke(self.title + part_title + info)
            else:
                context = self.full_retriever.invoke(self.title + part_title + info)

            context = self.limit_input_length(context)
                
            try:
                passage_content = generate_chain.invoke({
                    'context': context, 
                    'title': self.title, 
                    'abstract': info, 
                    'subpart_title': part_title,
                    'language': self.language,
                    'repeat_content': repeat_content,
                    'target_words': target_words + 100 if self.is_English else target_words - 100  # 添加字数要求
                    })
            except Exception as e:
                passage_content = backup_generate_chain.invoke({
                    'context': context, 
                    'title': self.title, 
                    'abstract': info, 
                    'subpart_title': part_title,
                    'language': self.language,
                    'repeat_content': repeat_content,
                    'target_words': target_words + 100 if self.is_English else target_words - 100  # 添加字数要求
                    })
                
            if self.log_path != '':
                self.write_log(content=passage_content, documents=context)

            passage_content = self.reduce_AIGC(passage_content)
            passage_content = clean_content(passage_content, is_English=self.is_English)
            subparts = seq_segment(passage_content)
            main_content = self.delete_seq_content(subparts)
            
            if self.word_count <= 5000: # 短文原生分段
                passage_ref_index = {}
                passage_contents = []
                for item in main_content:
                    if len(item) > 3:
                        item = self.find_refenerce(item, context, chapter_index, passage_ref_index)
                        passage_contents.append(item)
            else:
                passage_contents = ''.join(main_content)
                passage_contents = [self.find_refenerce(passage_contents, context, chapter_index)]
            print('单段落解析完成')
            print(passage_contents)

            return {'text': passage_contents, 'table': table, 'image': image}
        
        def add_multi_passages_with_context():
            # 添加表格
            table = None
            image = None
            if sub_part.title in self.select_chapter_with_table:
                table = self.__add_table(topic=info, table_type=self.table_type[0], is_analysis=True)
                image = self.__add_image(table=table)
            # -------------------------------段落总结------------------------------- # 
            if repeat_content == '':
                template = Prompt.multi_passage_template.pre_without_repeat
            else:
                template = Prompt.multi_passage_template.pre_with_repeat
                
            prompt = ChatPromptTemplate.from_template(template)
            chain = prompt | self.llm | StrOutputParser()
            backup_chain = prompt | OnlineModel(llm_name='GPT3.5').llm | StrOutputParser()
            
            context = self.full_retriever.invoke(self.title + part_title + info)
            try: 
                passages = chain.invoke({'context': context, 
                                    'title': self.title, 
                                    'abstract' : info, 
                                    'subpart_title': part_title,
                                    'passage_num': passage_num,
                                    'language': self.language,
                                    'form': Prompt.multi_passage_template.form,
                                    'repeat_content':repeat_content,
                                    # 'target_words': target_words,  # 添加总字数要求
                                    # 'target_words_per_passage': words_per_passage  # 添加每段字数要求
                                    })
            except Exception as e:
                passages = backup_chain.invoke({'context': context, 
                                    'title': self.title, 
                                    'abstract' : info, 
                                    'subpart_title': part_title,
                                    'passage_num': passage_num,
                                    'language': self.language,
                                    'form': Prompt.multi_passage_template.form,
                                    'repeat_content':repeat_content,
                                    # 'target_words': target_words,  # 添加总字数要求
                                    # 'target_words_per_passage': words_per_passage  # 添加每段字数要求
                                    })
            print('多段落总结生成完成')
            
            if self.log_path != '':
                self.write_log(content=passages, documents=context)
            
            try:
                match = re.search(r"(\{.*\})", passages, re.DOTALL)
                if match:
                    content = match.groups()[0]
                cc = content.strip()
                if cc.startswith('json```'):
                    cc = cc.replace('json```', '', 1)
                if cc.endswith('```'):
                    cc = cc[:-3]
                struct_data = json.loads(cc)
            except Exception as e:
                print(f"The response is not a valid JSON format: {e}")
                return None

            passages = []
            for key, item in struct_data.items():
                passages.append(item)
            print('多段落总结解析完成')

            # -------------------------------第一段------------------------------- #
            passage_contents = []
            passage_mian_content = passages[0]

            template = Prompt.multi_passage_template.first_passage
            prompt = ChatPromptTemplate.from_template(template)
            chain = prompt | self.llm | StrOutputParser()
            backup_chain = prompt | OnlineModel(llm_name='GPT3.5').llm | StrOutputParser()

            if sub_part.title in self.user_retrievers:
                context = self.local_retriever.invoke(self.title + part_title + passage_mian_content) + self.user_retrievers[sub_part.title].invoke(self.title + part_title + passage_mian_content)
            else:
                context = self.full_retriever.invoke(self.title + part_title + passage_mian_content)
            context = self.limit_input_length(context)

            try:
                passage_content = chain.invoke({'context':context, 
                    'passage_content': passage_mian_content,
                    'language': self.language,
                    'target_words_per_passage': words_per_passage + 100 if self.is_English else words_per_passage - 100 # 添加每段字数要求
                    })
            except Exception as e:
                passage_content = backup_chain.invoke({'context':context, 
                    'passage_content': passage_mian_content,
                    'language': self.language,
                    'target_words_per_passage': words_per_passage + 100 if self.is_English else words_per_passage - 100 # 添加每段字数要求
                    })
            
            print('多段落第一段解析完成')

            if self.log_path != '':
                self.write_log(content=passage_content, documents=context)

            passage_content = self.reduce_AIGC(passage_content)
            passage_content = clean_content(passage_content, is_English=self.is_English)
            subparts = seq_segment(passage_content)
            passage_content = ''.join(self.delete_seq_content(subparts))

            passage_content = self.find_refenerce(passage_content, context, chapter_index)
            passage_contents.append(passage_content)

            # -------------------------------后面段落------------------------------- #
            
            template = Prompt.multi_passage_template.subsequent_passage
            prompt = ChatPromptTemplate.from_template(template)
            chain = prompt | self.llm | StrOutputParser()
            backup_chain = prompt | OnlineModel(llm_name='GPT3.5').llm | StrOutputParser()

            before_passage = passage_content
            for passage_index in range(1, len(passages)):
                passage_mian_content = passages[passage_index]
                
                if sub_part.title in self.user_retrievers:
                    context = self.local_retriever.invoke(self.title + part_title + passage_mian_content) + self.user_retrievers[sub_part.title].invoke(self.title + part_title + passage_mian_content)
                else:
                    context = self.full_retriever.invoke(self.title + part_title + passage_mian_content)
                context = self.limit_input_length(context)

                try:
                    passage_content = chain.invoke({'context': context, 
                        'bafore_passage': before_passage,
                        'passage_content': passage_mian_content,
                        'language': self.language,
                        'target_words_per_passage': words_per_passage + 100 if self.is_English else words_per_passage - 100 # 添加每段字数要求
                        })
                except Exception as e:
                    passage_content = backup_chain.invoke({'context': context, 
                        'bafore_passage': before_passage,
                        'passage_content': passage_mian_content,
                        'language': self.language,
                        'target_words_per_passage': words_per_passage + 100 if self.is_English else words_per_passage - 100 # 添加每段字数要求
                        })
                    
                print('多段落第' + str(passage_index) +'段解析完成')

                if self.log_path != '':
                    self.write_log(content=passage_content, documents=context)

                passage_content = self.reduce_AIGC(passage_content)
                passage_content = clean_content(passage_content, is_English=self.is_English)
                subparts = seq_segment(passage_content)
                passage_content = ''.join(self.delete_seq_content(subparts))
                
                passage_content = self.find_refenerce(passage_content, context, chapter_index)
                passage_contents.append(passage_content)
                before_passage = passage_content

            return {'text': passage_contents, 'table': table, 'image': image}

        if passage_num == 1:
            passage_contents = add_single_passages()
        else:
            passage_contents = add_multi_passages_with_context()
                
        return passage_contents
    
    def __add_chapter(self, part: ChapterItem, chapter_index: int):  
        # 存在2、3级标题情况
        all_titles = []
        for sub_part in part.sub: # 遍历二级标题
            if len(sub_part.sub) == 0: # 二级标题没有三级标题
                all_titles.append(sub_part.title)
            else: # 二级标题有三级标题
                for subsub_part in sub_part.sub:
                    all_titles.append(subsub_part.title)

        part_content = {}

        for sub_part in part.sub:
            subpart_title = sub_part.title
            subsub_parts = sub_part.sub

            part_content[subpart_title] = {}
            
            if len(subsub_parts) == 0: 
                passage_num = random.randint(self.single_subsubpart_min_passage, self.single_subsubpart_max_passage)
                repeat_content = ''
                for item in all_titles:
                    if item != subpart_title:
                        repeat_content = repeat_content + '、' +  item
                passage_contents = self.__add_passages(sub_part, passage_num, chapter_index, repeat_content=repeat_content[1:])
                part_content[subpart_title]['only_one_part'] = passage_contents

            else:
                for subsub_part in subsub_parts:
                    passage_num = random.randint(self.multi_subsubpart_min_passage, self.multi_subsubpart_max_passage)
                    repeat_content = ''
                    for item in all_titles:
                        if item != subsub_part.title:
                            repeat_content = repeat_content + '、' +  item

                    passage_contents = self.__add_passages(subsub_part, passage_num, chapter_index, repeat_content=repeat_content[1:])
                    part_content[subpart_title][subsub_part.title] = passage_contents

        # 没有2、3级标题情况
        if len(part.sub) == 0:
            passage_num = random.randint(self.single_chapter_min_passage, self.single_chapter_max_passage)
            passage_contents = self.__add_passages(part, passage_num, chapter_index)
            part_content['only_one_part'] = passage_contents
        
        self.content_locker.lock.acquire()
        self.paper_content[part.title] = part_content
        self.content_locker.lock.release()

    def __add_image(self, table: str):
        # print(table)
        if table is None:
            return None
        url = "https://graphai.crazystone.work/foxchart/api/"
        data = {
            "data": table
        }
        headers = {
            'Authorization': 'sk-G3SpS1OgkfVBBXRVctpZBlbkzxoZJlokebqBJULE3fl9mt40'  # 替换为测试用户的API密钥
        }
        response = requests.post(url, json=data, headers=headers)
        if response.status_code == 200:
            timestamp_ms = int(time.time() * 1000)
            timestamp_ms_str = str(timestamp_ms)
            op = Global.tmp_path + timestamp_ms_str + '.png'
            with open(op, "wb") as f:
                f.write(response.content)
            return op
        else:
            print("生成图表失败:", response)
            return None

    def __add_table(self, topic: str, table_type: str, is_analysis: bool):
        # table_type包括三种类型：数据 理论 实验步骤
        prompt = ChatPromptTemplate.from_template(Prompt.table_template.create_table)
        chain = prompt | self.llm | StrOutputParser()
        backup_chain = prompt | OnlineModel(llm_name='GPT3.5').llm | StrOutputParser()
        try:
            table = chain.invoke({'topic': topic, 
                              'type': table_type, 
                              'form': Prompt.table_template.form, 
                              'language': self.language})
        except Exception as e:
            table = backup_chain.invoke({'topic': topic, 
                              'type': table_type, 
                              'form': Prompt.table_template.form, 
                              'language': self.language})
        print('表格生成完成')

        if '抱歉' in table or '无法' in table:
            return None
        
        try:
            match = re.search(r"(\{.*\})", table, re.DOTALL)
            if match:
                content = match.groups()[0]
            cc = content.strip()
            if cc.startswith('json```'):
                cc = cc.replace('json```', '', 1)
            if cc.endswith('```'):
                cc = cc[:-3]
            struct_data = json.loads(cc)
            print('表格解析完成')
        except Exception as e:
            print(f"The response is not a valid JSON format: {e}")
            return None
        
        analysis = None
        if is_analysis:
            prompt = ChatPromptTemplate.from_template(Prompt.table_template.analysis_table)
            chain = prompt | self.llm | StrOutputParser()
            backup_chain = prompt | OnlineModel(llm_name='GPT3.5').llm | StrOutputParser()
            try:
                row_analysis = chain.invoke({'table': table, 'topic': topic, 'language': self.language, 'words_per_section': self.words_per_section + 100 if self.is_English else self.words_per_section - 100}) 
            except Exception as e:
                row_analysis = backup_chain.invoke({'table': table, 'topic': topic, 'language': self.language, 'words_per_section': self.words_per_section + 100 if self.is_English else self.words_per_section - 100}) 
            row_analysis = self.reduce_AIGC(row_analysis)
            row_analysis = clean_content(row_analysis, is_English=self.is_English)

            main_content = self.delete_seq_content(seq_segment(row_analysis))
            if self.word_count <= 5000:
                analysis = main_content
            else:
                analysis = [''.join(main_content)]

        if self.log_path != '':
            self.write_log(content=table + row_analysis)
        print('表格整体完成')

        return {'table': struct_data, 'analysis': analysis}
    
    def __add_thanks(self):
        prompt = ChatPromptTemplate.from_template(Prompt.thanks_template.create_thanks)
        chain = prompt | self.llm | StrOutputParser()
        backup_chain = prompt | OnlineModel(llm_name='GPT3.5').llm | StrOutputParser()
        try:
            content = chain.invoke({# 'example': Prompt.thanks_template.example, 
                              'catagory': self.category,
                              'title': self.title,
                              'language': self.language,
                              'words_per_section': self.words_per_section + 100 if self.is_English else self.words_per_section - 100
                              })
        except Exception as e:
            content = backup_chain.invoke({# 'example': Prompt.thanks_template.example, 
                              'catagory': self.category,
                              'title': self.title,
                              'language': self.language,
                              'words_per_section': self.words_per_section + 100 if self.is_English else self.words_per_section - 100
                              })
        print('表格生成完成')
        content = self.reduce_AIGC(content)
        self.thanks = clean_content(re.sub('\n', '', content), is_English=self.is_English)
    
    # 搜索引用
    def find_refenerce(self, content:str, documents:list, chapter_index: int, passage_ref_index={}):
        def add_to_reference_list(refer_info: str) -> int:
            if refer_info in self.ref_papers:
                ref_index = self.ref_papers[refer_info]
            else:
                self.ref_locker.lock.acquire()
                ref_index = len(self.ref_papers)
                self.ref_papers[refer_info] = ref_index
                self.ref_locker.lock.release()
            return ref_index
        
        if chapter_index == len(self.outlines): # 最后一章不加引用
            return content
        
        if chapter_index == 1:
            ref_probability = 1
        elif chapter_index == 2:
            ref_probability = self.ref_second_chapter_prob
        else:
            ref_probability = self.ref_subsequent_chapter_prob
        
        add_refenerce_content = ''
        sentences = content.split('.' if self.is_English else '。')

        for sentence in sentences[:-1]:
            if len(sentence) == 0:
                continue

            if check_content(sentence, self.not_ref_text):
                add_refenerce_content += sentence
                add_refenerce_content += '.' if self.is_English else '。'
                continue

            papers = self.ref_databse.similarity_search_with_relevance_scores(query=sentence, k=self.ref_retrieve_num)
            ref_content = '['
            sentence_ref_index = []
            for i, paper in enumerate(papers):
                similarity = paper[1]
                refer_info = paper[0].metadata['reference']
                if random.random() < ref_probability and similarity >= self.ref_threshold:
                    ref_index = add_to_reference_list(refer_info)
                    
                    if ref_index not in passage_ref_index:
                        passage_ref_index[ref_index] = 1
                        if ref_index not in sentence_ref_index:
                            sentence_ref_index.append(ref_index)
                    elif random.random() < ref_probability * self.ref_continual_prob ** passage_ref_index[ref_index]:
                        passage_ref_index[ref_index] += 1
                        if ref_index not in sentence_ref_index:
                            sentence_ref_index.append(ref_index)

                if similarity < self.ref_threshold or len(sentence_ref_index) >= self.max_ref_num:
                    break
            
            if len(sentence_ref_index) == 0 and random.random() < ref_probability * self.ref_probability: # 最大相似度小于引用阈值，按概率进行引用
                refer_info = papers[0][0].metadata['reference']
                ref_index = add_to_reference_list(refer_info)

                if ref_index not in passage_ref_index:
                    passage_ref_index[ref_index] = 1
                    if ref_index not in sentence_ref_index:
                        sentence_ref_index.append(ref_index)
                elif random.random() < ref_probability * self.ref_continual_prob ** passage_ref_index[ref_index]:
                    passage_ref_index[ref_index] += 1
                    if ref_index not in sentence_ref_index:
                        sentence_ref_index.append(ref_index)

            for ref_index in sorted(sentence_ref_index):
                ref_content += str(ref_index) + ','   
            
            if len(ref_content) != 1:
                add_refenerce_content += sentence + ref_content[:-1] + ']'
            else:
                add_refenerce_content += sentence

            add_refenerce_content += '.' if self.is_English else '。'
        
        return add_refenerce_content

    # 删除序列词和总结词
    def delete_seq_content(self, content:list):
        is_delete = False
        if random.random() < self.delete_seq_prob:
            is_delete= True

        def merged_seq(begin: int, mark: str, row_content: str):
            seq_part = []
            for item in row_content:
                seq_part.append(True if item.find(mark) == begin - 1 else False)
            if sum(seq_part) == 0:
                return row_content

            temp = [row_content[0][begin:] if is_delete and seq_part[0] else row_content[0]]
            current_delete = True
            for index in range(1, len(seq_part)):
                if seq_part[index] and seq_part[index - 1]:
                    temp[-1] += row_content[index][begin:] if is_delete and current_delete else row_content[index] 
                elif seq_part[index] and not seq_part[index - 1] and len(row_content[index - 1]) < 50: # 前文是短句情况不删除
                    current_delete = False
                    temp[-1] += row_content[index] if temp[-1][-1] == '：' else '：' + row_content[index]
                else:
                    temp.append(row_content[index][begin:] if is_delete and seq_part[index] else row_content[index])

                if not seq_part[index] and seq_part[index - 1]:
                    current_delete = True
                
            return temp
        
        delete_content = merged_seq(begin=3, mark='，', row_content=content)
        delete_content = merged_seq(begin=2, mark= '.', row_content=delete_content)
        
        if delete_content[-1].find('综上所述') == 0:
            delete_content[-1] = random.choice(self.sum_word) + delete_content[-1][5:]
        return delete_content
    
    def write_log(self, content: str, documents=[]):
        context = []
        if len(documents) != 0:
            for item in documents:
                context.append((item.page_content, item.metadata))

        self.log_locker.lock.acquire()
        self.log_locker.content.append({'context': context, 'content': content})
        self.log_locker.lock.release()

    def WriteToDocx(self, output_path:str, forced_add_catalogs=False) -> None:
        # forced_add_catalogs 强制增加目录
        self.document = Document()
        # 标题
        self.__write_title()

        # 中文摘要
        self.__write_headline(headline='Abstract' if self.is_English else '摘要', level=0.5)
        self.__write_content(content=self.abstract[0])
        print('摘要写入完成')
        if self.is_English:
            try:
                self.__write_key_world(key_world='Keywords: ' +self.abstract[1])
            except Exception as e:
                print('关键词写入失败')
                print(self.abstract[1])
        else:
            try:
                self.__write_key_world(key_world='关键词：' +self.abstract[1])
            except Exception as e:
                print('关键词写入失败')
                print(self.abstract[1])
        self.document.add_page_break()
        print('关键词写入完成')

        # 英文摘要
        if self.EnglishAbstract != None:
            self.__write_headline(headline='Abstract', level=0.5)
            self.__write_content(content=self.EnglishAbstract[0], is_Chinese=False)
            self.__write_key_world(key_world='Keywords: ' + self.EnglishAbstract[1], is_Chinese=False)
            self.document.add_page_break()
            print('英文摘要写入完成')

        if forced_add_catalogs or self.word_count > 5000:
            self.__write_headline(headline='catalogs' if self.is_English else '目录', level=0.5)
            self.document.add_page_break()
            self.document.add_page_break()
            print('目录初始化写入完成')

        # 正文
        part_num = len(self.paper_content)
        seq_chapter = {}
        for current_part_num in range(part_num):
            str_num = str(current_part_num + 1) if self.is_English else self.number_to_chinese[current_part_num + 1]

            for i, (part_title, subpart_titles) in enumerate(self.paper_content.items()):
                print(part_title, '开始写入')
                if str_num in part_title:
                    seq_chapter[part_title] = subpart_titles

                    if self.word_count <= 5000:
                        self.__write_headline(headline=part_title.split(' ')[-1], level=1)
                    else:
                        self.__write_headline(headline=part_title, level=1)
  
                    if i == 0:
                        try:
                            self.generate_plantuml(self.paper_info.abstract)
                        except Exception as e:
                            print(f"生成 PlantUML 图像时出错: {e}")

                    for j, (subpart_title, subsubpart_titles) in enumerate(subpart_titles.items()):
                        print(subpart_title, '开始写入')
                        if subpart_title == 'only_one_part':
                            if random.random() < 0.5:  
                                # contents = "\n".join(str(item) for item in subsubpart_titles['text'])
                                contents = [content for content in subsubpart_titles['text']]
                                try:
                                    self.generate_plantuml(contents)
                                except Exception as e:
                                    print(f"生成 PlantUML 图像时出错: {e}")    
                            self.__random_insert_content(subsubpart_titles)
                           
                        else:
                            self.__write_headline(headline=subpart_title, level=2, level_format=(current_part_num + 1, j + 1))
                            for k, (subsubpart_title, passage_contents) in enumerate(subsubpart_titles.items()):
                                if subsubpart_title == 'only_one_part':
                                    if random.random() < 0.5: 
                                        # contents = "\n".join(str(item) for item in passage_contents['text'])
                                        contents = [content for content in passage_contents['text']]
                                        try:
                                            self.generate_plantuml(contents)
                                        except Exception as e:
                                            print(f"生成 PlantUML 图像时出错: {e}")    
                                    self.__random_insert_content(passage_contents)
                                else:
                                    self.__write_headline(headline=subsubpart_title, level=3, level_format=(current_part_num + 1, j + 1, k + 1))
                                    if random.random() < 0.5:  
                                        contents = [content for content in passage_contents['text']]
                                        try:
                                            self.generate_plantuml(contents)
                                        except Exception as e:
                                            print(f"生成 PlantUML 图像时出错: {e}")    
                                    self.__random_insert_content(passage_contents)

                    if self.word_count > 5000 or current_part_num == part_num - 1:
                        self.document.add_page_break()
                    break
        self.paper_content = seq_chapter
        
        # 参考文献
        self.__write_headline(headline='Reference' if self.is_English else '参考文献', level=1)
        self.__write_reference()
        self.document.add_page_break()
        print('参考文献写入完成')

        # 致谢
        if self.thanks != None: 
            self.__write_headline(headline='Acknowledgement' if self.is_English else '致谢', level=1)
            self.__write_content(content=self.thanks)
            print('致谢写入完成')

        # 目录更新
        if forced_add_catalogs or self.word_count > 5000:
            self.document.save(Global.tmp_path + self.title + '.docx')

            if self.EnglishAbstract != None:
                start = 12
            else:
                start = 7

            # 创建 Word 应用程序实例
            try:
                doc_app = win32.Dispatch('Word.Application')
                doc_app.Visible = True  # 可选，设置为 True 以查看 Word 窗口
            except Exception as e:
                print(f"无法启动 Word 应用程序: {e}")

            doc = doc_app.Documents.Add(Global.tmp_path + self.title + '.docx')

            parag_range = doc.Paragraphs(start).Range
            parag_range.InsertAfter('Catelogs')
            parag_range = doc.Paragraphs(start + 1).Range
            doc.TablesOfContents.Add(Range=parag_range, UseHeadingStyles=True,LowerHeadingLevel=3, UseHyperlinks=True)

            doc.SaveAs(output_path) 
            doc.Close() 
            doc_app.Quit()
            os.remove(Global.tmp_path + self.title + '.docx')
        else:
            self.document.save(output_path)

    def __write_title(self):
        part = self.document.add_heading(level=0)
        part.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = part.add_run(self.title)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0,0,0)
        run.font.name = 'Time New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def __write_headline(self, headline: str, level: int, level_format = None):
        if level == 0.5: # 摘要
            part = self.document.add_heading(level=1)
            run = part.add_run(headline)
            part.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run.font.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.name = 'Time New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif level == 1:
            part = self.document.add_heading(level=level)
            run = part.add_run(headline)
            part.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run.font.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.name = 'Time New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif level == 2:
            part = self.document.add_heading(level=level)
            format_headline = '{first}.{second} {headline}'.format(first=level_format[0], second=level_format[1], headline=headline)
            run = part.add_run(format_headline)
            run.font.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.name = 'Time New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif level == 3:
            part = self.document.add_heading(level=level)
            format_headline = '{first}.{second}.{third} {headline}'.format(first=level_format[0], second=level_format[1], third=level_format[2], headline=headline)
            run = part.add_run(format_headline)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0,0,0)
            run.font.name = 'Time New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def __random_insert_content(self, part: dict):
        if part is None:
            return
        if random.random() < 0.5: # 先写段落后写表格
            for index, passage_content in enumerate(part['text']):
                if index == len(part['text']) - 1 and part['table'] is not None and len(part['text']) != 1:
                    continue
        
                if len(passage_content) != 0:

                    print("===== 段落内容开始 =====")
                    print(passage_content)
                    print("===== 段落内容结束 =====")

                    self.__write_content(content=passage_content)

            if part['table'] is not None:
                try:
                    self.__write_table(part['table'])
                except:
                    print('Table wrong')

            if part['image'] is not None:
                try:
                    self.__write_image(part['image'])
                except:
                    print('Image wrong')
        else: # 先写表格后写段落
            if part is not None and part['table'] is not None:
                try:
                    self.__write_table(part['table'])
                except:
                    print('Table wrong')
            if part is not None and part['image'] is not None:
                try:
                    self.__write_image(part['image'])
                except:
                    print('Image wrong')

            for passage_content in part['text']:
                if len(passage_content) != 0:
                    print("===== 段落内容开始 =====")
                    print(passage_content)
                    print("===== 段落内容结束 =====")


                    self.__write_content(content=passage_content)

    def __write_content(self, content: str, is_Chinese = True):
        part = self.document.add_paragraph()
        part.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        part.paragraph_format.first_line_indent = Cm(0.8)
        part.paragraph_format.space_after = Pt(1) 

        run = part.add_run(content)
        run.font.size = Pt(12)
        run.font.name = 'Time New Roman'
        if is_Chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def __write_key_world(self, key_world: str, is_Chinese = True):
        part = self.document.add_paragraph()
        part = self.document.add_paragraph()
        part.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = part.add_run(key_world)
        run.font.size = Pt(12)
        run.font.name = 'Time New Roman'
        if is_Chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def __write_table(self, table_analysis: dict):
        def set_cell_border(cell, **kwargs):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # check for tag existnace, if none found, then create one
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)

            # list over all available tags
            for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
                edge_data = kwargs.get(edge)
                if edge_data:
                    tag = 'w:{}'.format(edge)

                    # check for tag existnace, if none found, then create one
                    element = tcBorders.find(qn(tag))
                    if element is None:
                        element = OxmlElement(tag)
                        tcBorders.append(element)

                    # looks like order of attributes is important
                    for key in ["sz", "val", "color", "space", "shadow"]:
                        if key in edge_data:
                            element.set(qn('w:{}'.format(key)), str(edge_data[key]))

        table = table_analysis['table'].get('Table')
        if table is None:
            return
        
        data = table # 用于生成图的数据

        table_list = []
        for row in table:
            column_list = row.split('|')
            table_list.append(column_list)
        col_num = len(column_list)
        row_num = len(table_list)

        table = self.document.add_table(rows=row_num, cols=col_num, style='Normal Table')

        # 向表格中添加数据
        for i in range(row_num):
            for j in range(col_num):
                table.cell(i, j).text = table_list[i][j]
                table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(12)
                table.cell(i, j).paragraphs[0].runs[0].font.name = 'Time New Roman'
                table.cell(i, j).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

                if i == 0:
                    set_cell_border(table.cell(i, j),
                                    top={"sz": 12, "val": "single", "color": "#000000"},
                                    bottom={"sz": 12, "val": "single", "color": "#000000"},
                                    # left={"sz": 12, "val": "single", "color": "#FFFFFF"},
                                    # right={"sz": 12, "val": "single", "color": "#FFFFFF"},
                                    )
                if i == row_num - 1:
                    set_cell_border(table.cell(i, j),
                                    bottom={"sz": 12, "val": "single", "color": "#000000"},
                                    )

        # 获取大模型生成的 matplotlib 代码
        from langchain.prompts import ChatPromptTemplate
        from langchain.schema.output_parser import StrOutputParser
        import matplotlib.pyplot as plt
        import numpy as np
        import os

        # 构建 prompt
        template = """请根据以下表格内容生成一段 matplotlib 的代码，要求生成一个合适的能够用于论文的数据图，不要plt.show()和plt.savefig()，字体设置为'SimHei'：
        {data}

        如果表格内容没有数据，请根据情景生成看起来合理的数据用于论文绘图。
        """
        prompt = ChatPromptTemplate.from_template(template)
        chain = prompt | self.llm | StrOutputParser()

        # 获取大模型生成的 matplotlib 代码
        matplotlib_code = chain.invoke({"data": data})

        # 使用正则表达式提取代码块
        code_match = re.search(r'```python(.*?)```', matplotlib_code, re.DOTALL)
        if code_match:
            extracted_code = code_match.group(1).strip()  # 提取代码并去除多余空白
            # print("提取的代码：", extracted_code)  # 打印提取的代码以供调试

            # 执行提取的 matplotlib 代码
            try:
                exec(extracted_code)
            except Exception as e:
                print(f"执行 matplotlib 代码时出错: {e}")
        else:
            print("未找到有效的代码块。")

        # 保存图像
        image_path = Global.tmp_path + str(int(time.time() * 1000)) + '.png'
        plt.tight_layout()
        try:
            plt.savefig(image_path)
            plt.close()
        except Exception as e:
            print(f"保存图像时出错: {e}")

        # 将图像插入到现有的 Word 文档中
        from docx import Document
        from docx.shared import Inches
        try:
            self.document.add_picture(image_path, width=Inches(5.0))  # 根据需要调整图像大小
        except Exception as e:
            print(f"插入图像到 Word 文档时出错: {e}")

        # 清理生成的图像文件
        if os.path.exists(image_path):
            try:
                os.remove(image_path)
            except Exception as e:
                print(f"清理图像文件时出错: {e}")

        # 添加表格的分析部分
        analysis = table_analysis['analysis']
        if analysis is not None:
            for item in analysis:
                self.__write_content(content=item)

    def __write_image(self, image_file: str):
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = para.add_run("")
        rr.add_picture(image_file, width=Inches(4))

    def __write_reference(self):
        ref_format = '[{num}] {reference}\n'
        del self.ref_papers['0']

        part = self.document.add_paragraph()

        for reference, num in self.ref_papers.items():
            run = part.add_run(ref_format.format(num=num, reference=reference))
            run.font.size = Pt(12)

            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def generate_plantuml(self, outline_str):
        # 将大纲转换为字符串格式
        # def parse_chapter(chapter):
        #     # 生成章节标题
        #     chapter_str = f"* {chapter.title}\n"
        #     # 递归解析子章节
        #     for sub in chapter.sub:
        #         chapter_str += parse_chapter(sub)
        #     return chapter_str

        # outline_str = "\n".join([parse_chapter(chapter) for chapter in self.paper_info.outline])
        
        
        # 打印大纲以供调试
        # print("生成的大纲：", outline_str)
        
        # 使用 LangChain 获取 PlantUML 代码
        prompt = ChatPromptTemplate.from_template("请根据以下内容生成能够合理描述其内容的架构设计图或内容描述图的 PlantUML 代码：{outline}")
        chain = prompt | self.llm | StrOutputParser()
        plantuml_response = chain.invoke({"outline": outline_str})
        
        # 提取 PlantUML 代码
        plantuml_code_match = re.search(r'```plantuml(.*?)```', plantuml_response, re.DOTALL)
        if plantuml_code_match:
            plantuml_code = plantuml_code_match.group(1).strip()  # 提取代码并去除多余空白
            # print("提取的 PlantUML 代码：", plantuml_code)  # 可选，打印提取的代码以供调试

            # 使用 PlantUML 生成图像
            plantuml_url = "http://www.plantuml.com/plantuml/png/"
            plantuml = PlantUML(url=plantuml_url)

            try:
                # 处理 PlantUML 代码
                png_data = plantuml.processes(plantuml_code)
                
                # 保存 PNG 文件
                image_path = Global.tmp_path + str(int(time.time() * 1000)) + '.png'  # 使用时间戳生成文件名
                with open(image_path, 'wb') as f:
                    f.write(png_data)

                # 将图像插入到 Word 文档中
                self.document.add_picture(image_path, width=Inches(5.0))  # 根据需要调整图像大小

            except Exception as e:
                print(f"生成 PlantUML 图像时出错: {e}")

            # 清理生成的图像文件
            if os.path.exists(image_path):
                os.remove(image_path)
        else:
            print("未找到有效的 PlantUML 代码块。")

class OpeningReportWriter(BaseWriter):
    def __init__(self, llm: OnlineModel, paper_info: PaperInfoAll, paper_writer: PaperWriter) -> None:
        super().__init__(llm, paper_info, '')
        self.paper_content = paper_writer.paper_content
        self.ref_papers = paper_writer.ref_papers

        self.template_path = 'D:\WordTemplate\OpeningReport.docx'

    def WriteOpeningReport(self, output_path: str, span_time = '1'):
        self.document = Document(self.template_path)

        self.__write_title(self.title)

        # 背景
        background = {}
        key = None
        for key in self.paper_content: break
        if key is not None:
            for subpart_title, subsubpart_titles in self.paper_content[key].items():
                if subpart_title == 'only_one_part':
                    background[key.split(' ')[-1]] = subsubpart_titles['text']
                else:
                    for subsubpart_title, passage_contents in subsubpart_titles.items():
                        if subsubpart_title == 'only_one_part':
                            background[subpart_title] = passage_contents['text']
                        else:
                            background[subsubpart_title] = passage_contents['text']
        
        self.__write_headline('一、立题依据', level=1)
        for index, (title, content) in enumerate(background.items()):
            self.__write_headline(str(index + 1) + '、' + title, level=2)
            self.__write_content(content)

        # 方法
        self.__write_headline('二、技术路线', level=1)
        for i in range(1, len(self.outlines)):
            self.__write_headline(str(i) + '、' + self.outlines[i].title.split(' ')[-1], level=2)
            self.__write_content(self.outlines[i].info)
            for j, subpart in enumerate(self.outlines[i].sub):
                self.__write_content('（' + str(j + 1) + '）' + subpart.title)
                self.__write_content(subpart.info)

        # 参考文献
        self.__write_headline('三、参考文献', level=1)
        self.__write_reference()

        # 工作进度安排
        self.__write_headline('四、工作进度安排', level=1)
        current_time = datetime.datetime.now()
        current = str(current_time.year) + '年' + str(current_time.month) + '月' + str(current_time.day) + '日'
        prompt = ChatPromptTemplate.from_template(Prompt.research_plan.template)
        chain = prompt | self.llm | StrOutputParser()
        backup_chain = prompt | OnlineModel(llm_name='GPT3.5').llm | StrOutputParser()
        try:
            content = chain.invoke({'catagory': self.category,
                                'abstract': self.abstract[0],
                                'title': self.title,
                                'begin_time': current,
                                'span_time': span_time,
                              })
        except Exception as e:
            content = backup_chain.invoke({'catagory': self.category,
                                'abstract': self.abstract[0],
                                'title': self.title,
                                'begin_time': current,
                                'span_time': span_time,
                              })
        self.plan = content
        for item in seq_segment(content):
            self.__write_content(item)

        # 指导老师对开题报告的意见
        self.__write_headline('五、指导老师对开题报告的意见', level=1)

        self.__write_sign()
        self.document.save(output_path)

    def __write_title(self, content: str, is_Chinese = True):
        for para in self.document.paragraphs:
            if '论文题目' in para.text:
                run = para.add_run(content)
                run.font.size = Pt(18)
                run.font.name = 'Time New Roman'
                if is_Chinese:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def __write_content(self, content: str, is_Chinese = True):
        part = self.document.tables[0].cell(0, 0).add_paragraph()
        part.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        part.paragraph_format.first_line_indent = Cm(0.8)
        part.paragraph_format.space_after = Pt(1) 

        run = part.add_run(content)
        run.font.size = Pt(12)
        run.font.name = 'Time New Roman'
        if is_Chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def __write_headline(self, content: str, level: int, is_Chinese = True):
        part = self.document.tables[0].cell(0, 0).add_paragraph()
        part.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = part.add_run(content)
        run.font.size = Pt(16) if level == 1 else Pt(14)
        run.font.name = 'Time New Roman'
        if is_Chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def __write_sign(self):
        for i in range(3):
            part = self.document.tables[0].cell(0, 0).add_paragraph()
            run = part.add_run('')
            run.font.size = Pt(12)
            run.font.name = 'Time New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        part = self.document.tables[0].cell(0, 0).add_paragraph()
        run = part.add_run('                                                 指导老师签名：')
        run.font.size = Pt(12)
        run.font.name = 'Time New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        part = self.document.tables[0].cell(0, 0).add_paragraph()
        run = part.add_run('                                                     年     月     日')
        run.font.size = Pt(12)
        run.font.name = 'Time New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def __write_reference(self):
        ref_format = '[{num}] {reference}\n'
        part = self.document.tables[0].cell(0, 0).add_paragraph()

        for reference, num in self.ref_papers.items():
            run = part.add_run(ref_format.format(num=num, reference=reference))
            run.font.size = Pt(12)

            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

class TaskBookWriter(BaseWriter):
    def __init__(self, llm: OnlineModel, paper_info: PaperInfoAll, paper_writer: PaperWriter, opening_writer: OpeningReportWriter) -> None:
        super().__init__(llm, paper_info, '')
        self.paper_content = paper_writer.paper_content
        self.ref_papers = paper_writer.ref_papers

        if opening_writer is None:
            self.plan = None
        else:
            self.plan = opening_writer.plan

        self.template_path = 'D:\WordTemplate\TaskBook.docx'

    def WriteTaskBook(self, output_path: str, span_time = '1'):
        chain = self.llm | StrOutputParser()

        background = ''
        for key in self.paper_content: break
        for subpart_title, subsubpart_titles in self.paper_content[key].items():
            if subpart_title == 'only_one_part':
                background += '***' + key.split(' ')[-1] + '***'
                for item in subsubpart_titles['text']:
                    background += item
            else:
                for subsubpart_title, passage_contents in subsubpart_titles.items():
                    if subsubpart_title == 'only_one_part':
                        background += '***' + subpart_title + '***'
                        for item in passage_contents['text']:
                            background += item
                    else:
                        background += '***' + subsubpart_title + '***'
                        for item in passage_contents['text']:
                            background += item

        headlines = ''
        for i in range(1, len(self.outlines)):
            if len(self.outlines[i].sub) == 0:
                headlines += self.outlines[i].title.split(' ')[-1] + '、'
            else:
                for j, subpart in enumerate(self.outlines[i].sub):
                    headlines += subpart.title + '、'
        
        if self.plan == None:
            current_time = datetime.datetime.now()
            current = str(current_time.year) + '年' + str(current_time.month) + '月' + str(current_time.day) + '日'
            title_org, exp_task, plan = chain.batch([Prompt.task_book.tilte_org.format(title=self.title, background=background),
                                                Prompt.task_book.exp_task.format(title=self.title, abstract=self.abstract[0] ,subtitles=headlines[:-1] + '。'),
                                                Prompt.research_plan.template.format(catagory=self.category, abstract=self.abstract[0], title=self.title, begin_time=current, span_time=span_time)])
        else:
             plan = self.plan
             title_org, exp_task = chain.batch([Prompt.task_book.tilte_org.format(title=self.title, background=background),
                                                Prompt.task_book.exp_task.format(title=self.title, abstract=self.abstract[0] ,subtitles=headlines[:-1] + '。')])

        
        self.document = Document(self.template_path)

        self.__write_title(self.title)
        # 题目来源
        self.__write_headline('一、题目来源', level=1)
        self.__write_content(clean_content(re.sub('\n', '', title_org)))

        # 基本任务要求
        self.__write_headline('二、基本任务要求', level=1)
        self.__write_headline('基本实验任务包括以下几点：', level=2)
        for item in seq_segment(exp_task):
            self.__write_content(item)

        self.__write_headline('基本论文写作任务包括以下几点：', level=2)
        for i in range(len(self.outlines)):
            self.__write_content(str(i + 1) + '、' + self.outlines[i].title.split(' ')[-1])
            self.__write_content(self.outlines[i].info)
            for j, subpart in enumerate(self.outlines[i].sub):
                self.__write_content('（' + str(j + 1) + '）' + subpart.title + '：' + subpart.info)

        # 进度安排
        self.__write_headline('三、工作进度安排', level=1)
        for item in seq_segment(plan):
            self.__write_content(item)

        # 主要参考文献
        self.__write_headline('四、参考文献', level=1)
        self.__write_reference()
        
        # 其他要求
        self.__write_headline('五、其他要求', level=1)
        self.__write_content('无')

        self.document.save(output_path)

    def __write_title(self, content: str, is_Chinese = True):
        for para in self.document.paragraphs:
            if '论文题目' in para.text:
                run = para.add_run(content)
                run.font.size = Pt(18)
                run.font.name = 'Time New Roman'
                if is_Chinese:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def __write_content(self, content: str, is_Chinese = True):
        part = self.document.tables[0].cell(0, 0).add_paragraph()
        part.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        part.paragraph_format.first_line_indent = Cm(0.8)
        part.paragraph_format.space_after = Pt(1) 

        run = part.add_run(content)
        run.font.size = Pt(12)
        run.font.name = 'Time New Roman'
        if is_Chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def __write_headline(self, content: str, level: int, is_Chinese = True):
        part = self.document.tables[0].cell(0, 0).add_paragraph()
        part.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run = part.add_run(content)
        run.font.size = Pt(16) if level == 1 else Pt(14)
        run.font.name = 'Time New Roman'
        if is_Chinese:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def __write_reference(self):
        ref_format = '[{num}] {reference}\n'
        part = self.document.tables[0].cell(0, 0).add_paragraph()

        for reference, num in self.ref_papers.items():
            run = part.add_run(ref_format.format(num=num, reference=reference))
            run.font.size = Pt(12)

            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
