from LargeModel import OnlineModel
from Interface import BaseWriter, PaperInfoAll, ChapterItem
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import json
import Prompt
import torch

class SurveyWriter(BaseWriter):
    def __init__(self, llm: OnlineModel, paper_info: PaperInfoAll):
        super().__init__(llm, paper_info, '')

        
    def WriteSurvey(self, output_path: str):
        """生成调查问卷"""
        self.document = Document()
        
        # 设置页边距
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)
            
        # 使用大模型生成标题和前言
        header_prompt = Prompt.survey_template.header.format(
            title=self.title,
            category=self.category
        )
        # print("\n=== 标题和前言的 Prompt ===")
        # print(header_prompt)
        
        header_content = self.llm.invoke(header_prompt).content
        # print("\n=== 标题和前言的响应 ===")
        # print(header_content)
        
        # 清理标题和前言的响应
        header_content = header_content.replace('**', '').replace("#", "").replace('调查问卷标题：', '').replace('前言：', '').strip()
        header_parts = header_content.split('\n', 1)
        
        # 确保标题和前言都被正确分离
        if len(header_parts) == 2:
            title = header_parts[0].strip()
            introduction = header_parts[1].strip()
            self.__write_title(title)
            self.__write_introduction(introduction)
        else:
            raise ValueError("标题和前言格式不正确")
        
        # 使用大模型生成问题
        questions_prompt = Prompt.survey_template.questions.format(
            title=self.title,
            abstract=self.abstract[0],
            category=self.category
        )
        # print("\n=== 问题生成的 Prompt ===")
        # print(questions_prompt)
        
        response = self.llm.invoke(questions_prompt).content
        # print("\n=== 问题生成的原始响应 ===")
        # print(response)
        # print("\n=== 响应类型 ===")
        # print(type(response))
        
        try:
            # 尝试清理响应文本
            response = response.strip()
            # 移除可能存在的星号
            response = response.replace('**', '')
            # print("\n=== 清理前缀后缀前的响应 ===")
            # print(response)
            
            if response.startswith('```'):
                response = response.replace('```json\n', '').replace('```', '')
            
            # print("\n=== 清理后的响应 ===")
            # print(response)
            
            # 确保响应是一个完整的 JSON 对象
            if not response.startswith('{'):
                response = '{' + response
            if not response.endswith('}'):
                response = response + '}'
            
            # print("\n=== 最终处理的响应 ===")
            # print(response)
            
            questions_content = json.loads(response)
            # print("\n=== 解析后的 JSON 内容 ===")
            # print(json.dumps(questions_content, indent=2, ensure_ascii=False))
            
            # 验证JSON结构
            required_keys = ["single_choice", "multiple_choice", "open_questions"]
            for key in required_keys:
                if key not in questions_content:
                    raise ValueError(f"Missing required key: {key}")
            
            # 写入单选题
            for i, q in enumerate(questions_content["single_choice"], 1):
                self.__write_single_choice(i, q["question"], q["options"])
            
            # 写入多选题
            current_num = len(questions_content["single_choice"]) + 1
            for q in questions_content["multiple_choice"]:
                self.__write_multiple_choice(current_num, q["question"], q["options"])
                current_num += 1
            
            # 写入开放性问题
            for q in questions_content["open_questions"]:
                self.__write_fill_blank(current_num, q)
                current_num += 1
            
        except json.JSONDecodeError as e:
            print(f"JSON解析错误: {e}")
            print(f"原始响应: {response}")
            raise
        except Exception as e:
            print(f"处理问题时发生错误: {e}")
            raise
        
        # 写入结束语
        self.__write_ending("再次感谢您的支持与配合，您的回答对我们的研究意义重大。")
        
        # 保存文档
        self.document.save(output_path)
        print(f"问卷已生成：{output_path}")
        
    def __write_title(self, content: str):
        """写入标题"""
        title = self.document.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run(content)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
    def __write_introduction(self, content: str):
        """写入前言"""
        intro = self.document.add_paragraph()
        intro.paragraph_format.first_line_indent = Pt(24)
        run = intro.add_run(content)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
    def __write_single_choice(self, num: int, question: str, options: list):
        """写入单选题"""
        q = self.document.add_paragraph()
        q.paragraph_format.first_line_indent = Pt(24)
        run = q.add_run(f"{num}. {question}（单选）\n")
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 写入选项
        option_text = ""
        for i, opt in enumerate(options):
            option_text += f"{chr(65+i)}. {opt} "
        run = q.add_run(option_text + "\n")
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
    def __write_multiple_choice(self, num: int, question: str, options: list):
        """写入多选题"""
        q = self.document.add_paragraph()
        q.paragraph_format.first_line_indent = Pt(24)
        run = q.add_run(f"{num}. {question}（多选）\n")
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 写入选项
        option_text = ""
        for i, opt in enumerate(options):
            option_text += f"{chr(65+i)}. {opt} "
        run = q.add_run(option_text + "\n")
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
    def __write_fill_blank(self, num: int, question: str):
        """写入填空题"""
        q = self.document.add_paragraph()
        q.paragraph_format.first_line_indent = Pt(24)
        run = q.add_run(f"{num}. {question}（填空）\n")
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 添加填空线
        run = q.add_run("_" * 50 + "\n")
        run.font.size = Pt(12)
        
    def __write_ending(self, content: str):
        """写入结束语"""
        ending = self.document.add_paragraph()
        ending.paragraph_format.first_line_indent = Pt(24)
        run = ending.add_run(content)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def main():
    try:
        # 加载摘要
        ChineseAbstract = torch.load('C1')
        # print("摘要内容:", ChineseAbstract)
        
        # 创建 paper_info
        paper_info = PaperInfoAll(
            category='工商管理', 
            title='劳动促进经济的发展', 
            wordCount=10000, 
            useThreeLevel=False, 
            needGenOpeningReport=True,
            needGenTaskBook=True,
            needGenPPT=True,
            abstract=ChineseAbstract, 
            outline=[]  # 问卷不需要大纲
        )
        
        # 创建大模型实例
        MyModel = OnlineModel(llm_name='GPT3.5')
        
        # 创建问卷生成器实例
        survey = SurveyWriter(MyModel, paper_info)
        
        # 生成问卷
        output_dir = 'D:\\release\\Survey'
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, '调查问卷.docx')
        
        survey.WriteSurvey(output_path)
        # print(f"问卷已生成：{output_path}")
        
    except Exception as e:
        import traceback
        print("\n=== 错误详情 ===")
        print(f"错误类型: {type(e).__name__}")
        print(f"错误信息: {str(e)}")
        print("\n=== 完整的错误堆栈 ===")
        traceback.print_exc()

if __name__ == "__main__":
    main()