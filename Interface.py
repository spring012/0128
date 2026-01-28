import Prompt
import threading

from abc import ABC, abstractmethod
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
import threading
from pydantic import BaseModel
from typing import List, ForwardRef, Tuple
from oss import OSS
from docx import Document
import time
import Global
import os,re,random
from LargeModel import LargeLanguageModel,OnlineModel
from langdetect import detect

MAX_PROMPT_LENGTH = 4096 - 300
MAX_CONTEXT_LENGTH= 16385
class Locker:
    lock = threading.Lock()
    content = []

ChapterInfoModel = ForwardRef('ChapterItem')
class ChapterItem(BaseModel):
    title: str = '',
    info: str = '',
    referenceFileList: List[str] = [],
    sub: List[ChapterInfoModel]

class PaperInfo(BaseModel):
    model: str = '3.5'  # 可选值：3.0或4.0
    isEnglish: bool = False  # 是否是生成英文论文
    category: str = '0401教育学类'  # 科目
    title: str = '新时代中小学教师转型'  # 论文题目
    wordCount: int  # 论文字数
    useThreeLevel: bool = False  # 是否使用三级大纲
    needAIGC: bool = False # 是否需要降低AIGC率
    customOutline: str = '' # 用户自定义的大纲
    language: str = "zh"
    def __init__(self, **data):
        super().__init__(**data)
        try:
            lang = detect(self.title)
            self.language = lang
            self.isEnglish = lang.startswith("en")  # 如果标题是英文，就设置为 True
        except Exception as e:
            self.language = 'zh'
            self.isEnglish = False

class PaperInfoAll(BaseModel):
    orderid: str = '' # 内部订单id
    model: str = '3.5'  # 可选值：3.0或4.0
    isEnglish: bool = False  # 是否是生成英文论文
    category: str = '0401教育学类'  # 科目
    title: str = '新时代中小学教师转型'  # 论文题目
    wordCount: int = 10000 # 论文字数
    useThreeLevel: bool = False  # 是否使用三级大纲
    needGenOpeningReport: bool = False # 是否生成开题报告
    needGenTaskBook: bool = False # 是否生成任务书
    needGenPPT: bool = False # 是否生成答辩ppt
    aiPercent: int = 5 # 0-10之间，使用用户投喂数据数量，0是一个都不用，10是全用用户数据
    needAIGC: bool = False # 是否需要降低AIGC率
    abstract: List = []
    referenceFileList: List = []
    outline: List[ChapterItem] = []
    needSurvey: bool = False # 是否生成调查问卷

class AIGCContext(BaseModel):
    text: str = '' # 需要降低AIGC率的文本内容

class AIGCContextV2(BaseModel):
    text: str = '' # 需要降低AIGC率的文本内容
    lang: str = 'zh' # 文件语言类型：支持zh, en

class AIGCFileContext(BaseModel):
    fileurl: str = '' # 需要降低AIGC率的文件url，支持docx，txt
    orderid: str = '' # 内部订单id

class AIGCFileContextV2(BaseModel):
    fileurl: str = '' # 需要降低AIGC率的文件url，支持docx，txt
    orderid: str = '' # 内部订单id
    lang: str = 'zh' # 文件语言类型：支持zh, en

class DataBase(ABC):
    def __init__(self, model_name: str, data_path: str, database_path: str, device: str, **kwargs) -> None:
        """
        model_name: 检索器使用模型
        data_path: 存放pdf的文件夹路径
        database_path: 向量数据库路径
        device: 选择模型运行的设备cuda或者cpu
        作用: 初始化或加载向量数据库
        """
        self.__embedding_model = self.__load_model(model_name)
        self.embedding_database = self.__init_database(data_path, database_path)


    @abstractmethod
    def load_model(self, model_name:str, **kwargs):
        """
        model_name: 检索器使用模型
        作用: 加载数据库编码模型
        """
        pass
    
    @abstractmethod
    def init_database(self, data_path:str, database_path:str, **kwargs):
        """
        data_path: 存放pdf的文件夹路径
        database_path: 向量数据库路径
        作用: 初始化数据库
        """
        pass
    
    def retrieve(self, query:str, tok_k=15, **kwargs) -> list:
        """
        query: 查询语句
        tok_k: 返回结果数量
        Return: 与输入查询相关的片段列表
        作用: 检索数据库中与query相关的tok_k个样本
        """
        docs = self.embedding_database.similarity_search(query, k=tok_k)
        return docs

class RetrieverFramework(ABC):
    def __init__(self, ) -> None:
        pass

class BaseWriter():
    def __init__(self, llm: LargeLanguageModel, paper_info: PaperInfoAll, log_path=''):
        self.llm = llm.llm
        self.llm_abstract = llm.llm_abstract
        self.paper_info = paper_info

        self.title = paper_info.title
        self.category = paper_info.category
        self.log_path = log_path
        if paper_info.wordCount == 4500 or paper_info.wordCount == 4800:
            self.word_count = 10000
        else:
            self.word_count = paper_info.wordCount
        self.is_English = paper_info.isEnglish
        self.language = 'English' if self.is_English else 'Chinese'
        self.needAIGC = paper_info.needAIGC

        self.aigc_locker = threading.Lock()

        if type(paper_info) is PaperInfoAll:
            self.abstract = paper_info.abstract
            self.outlines = paper_info.outline
            
    
    def limit_input_length(self, context: list):
        for i in range(len(context)):
            context[i].page_content = context[i].page_content[:1024]
        return context
    
    def reduce_AIGC(self, context: str):
        language_config = {
        "Chinese":{
            "dot":'。',
            "comma":'，'
        },
        "English":{
            "dot":'.',
            "comma":',',
        }}
        content = context
        if self.needAIGC:
            self.llm = OnlineModel(llm_name='Zhipu-'+self.language).llm
            chain =  ChatPromptTemplate.from_template(Prompt.reduce_AIGC_template.English_prompt if self.is_English else Prompt.reduce_AIGC_template.Chinese_prompt) | self.llm | StrOutputParser()
            sent = context.split(language_config[self.language]["dot"])
            data = []
            for item in sent:
                if len(item) > 1:
                    data.append({'context': item + language_config[self.language]["dot"]})
            
            self.aigc_locker.acquire()
            result = []
            if len(data) != 0:
                result = chain.batch(data)
            self.aigc_locker.release()
            
            content = ''
            for item in result:
                # 使用正则表达式分割句子为最小片段（没有标点符号的部分）
                parts = re.split('['+language_config[self.language]["dot"]+language_config[self.language]["comma"]+']', item)
                parts = [part.strip() for part in parts if part.strip()]  # 去掉空白部分
                
                # 随机乱序
                random.shuffle(parts)
                
                # 重组句子
                if parts:  # 确保分割后有内容
                    shuffled_sentence = language_config[self.language]["comma"].join(parts) + language_config[self.language]["dot"]  # 中间部分用逗号，末尾用句号
                    content += shuffled_sentence
        
        return content
    
class AIGCTool():
    def __init__(self, llm: LargeLanguageModel):
        self.llm = llm.llm
        self.aigc_locker = threading.Lock()
        self.is_English = False
        self.language = 'English' if self.is_English else 'Chinese'
    def __init__(self, llm: LargeLanguageModel, en: bool):
        self.llm = llm.llm
        self.aigc_locker = threading.Lock()
        self.is_English = en
        self.language = 'English' if self.is_English else 'Chinese'
    def reduceAIGC(self, context: str):
        language_config = {
        "Chinese":{
            "dot":'。',
            "comma":'，'
        },
        "English":{
            "dot":'.',
            "comma":',',
        }}
        self.llm = OnlineModel(llm_name='Doubao-'+self.language).llm
        chain =  ChatPromptTemplate.from_template(Prompt.reduce_AIGC_template.English_prompt if self.is_English else Prompt.reduce_AIGC_template.Chinese_prompt) | self.llm | StrOutputParser()
        sent = context.split(language_config[self.language]["dot"])
        data = []
        for item in sent:
            if len(item) > 1:
                data.append({'context': item + language_config[self.language]["dot"]})
        
        self.aigc_locker.acquire()
        result = []
        if len(data) != 0:
            result = chain.batch(data)
        self.aigc_locker.release()
        
        content = ''
        for item in result:
            # 确保 item 是字符串
            if isinstance(item, str):
                parts = re.split('[' + language_config[self.language]["dot"] + language_config[self.language]["comma"] + ']', item)
                parts = [part.strip() for part in parts if part.strip()]  # 去掉空白部分

                # 取消随机乱序
                # random.shuffle(parts)

                # 重组句子
                if parts:  # 确保分割后有内容
                    shuffled_sentence = language_config[self.language]["comma"].join(parts) + language_config[self.language]["dot"]  # 中间部分用逗号，末尾用句号
                    content += shuffled_sentence
                else:
                    content += item
            else:
                print(f"Warning: item {item} is not a string and will be skipped.")

        return content

    def reduceDocxAIGC(self, docx_file_path: str) -> str:
        filename = os.path.basename(docx_file_path)  # Extract filename from docx_file_path
        doc = Document(docx_file_path)
        results = []
        for para in doc.paragraphs:
            reduced_text = self.reduceAIGC(para.text)
            results.append(reduced_text)
        new_doc = Document()
        for para in results:
            new_doc.add_paragraph(para)
        timestamp_ms = int(time.time() * 1000)
        timestamp_ms_str = str(timestamp_ms)
        output_path = Global.paper_save_path +timestamp_ms_str + '_' + filename
        new_doc.save(output_path)
        urls = OSS().upload(filename, output_path)
        return urls
    def reduceTxtAIGC(self, txt_file_path: str) -> str:
        filename = os.path.basename(txt_file_path)  # Extract filename from txt_file_path
        with open(txt_file_path, 'r', encoding='utf-8') as file:
            content = file.read()  # Read the content of the txt file
        reduced_content = self.reduceAIGC(content)  # Process the content
        timestamp_ms = int(time.time() * 1000)
        timestamp_ms_str = str(timestamp_ms)
        output_path = Global.paper_save_path + timestamp_ms_str + '_' + filename  # Define output path
        with open(output_path, 'w', encoding='utf-8') as file:
            file.write(reduced_content)  # Write the reduced content to the output path
        urls = OSS().upload(filename, output_path)  # Upload the file and get URLs
        return urls
    
class PaperWriterFramework(ABC):
    def __init__(self, pdf_db:DataBase, llm:LargeLanguageModel, title:str, category:str, retrieve_num:int, locker:Locker, save_log=False ,**kwargs) -> None:
        """
        pdf_db: PDF向量数据库
        llm: 生成论文使用的大模型
        title: 论文题目
        category: 论文学科
        retrieve_num: 生成正文时检索文章数目
        locker: 章节间并行运行的锁
        save_log: 是否保存日志文件
        """
        self.embedding_database = pdf_db.embedding_database
        self.llm_pipline = llm.llm_pipline

    @abstractmethod
    def GetAbstract(self, min_abstract_word:int, **kwargs):
        """
        min_abstract_word: 摘要最小字数
        Return: 文章的摘要和关键字
        作用: 写出文章的摘要
        """
        pass

    @abstractmethod
    def GetEnglishAbstract(self, **kwargs) -> list:
        """
        作用: 写出文章的英文摘要
        """
        pass
    
    @abstractmethod
    def GetOutline(self, UseTreeLevel:bool, **kwargs) -> list:
        """
        useThreeLevel: 是否使用三级大纲，表示生成大纲的最大层级
        Return: 文章的提纲
        作用: 写出文章的摘要
        """
        pass

    def ContentModify(self, new_ChineseAbstract:tuple, new_outlines:str) -> None:
        """
        new_ChineseAbstract: 修改后的摘要
        new_outlines: 修改后的提纲
        Return: None
        作用: 向类内传入用户修改后的提纲和摘要
        """
        self.ChineseAbstract = new_ChineseAbstract
        self.outlines = new_outlines

    @abstractmethod
    def GetFullPaper(self, **kwargs) -> dict:
        """
        Return: dict形式的文章，key为文章标题，value为标题内容
        作用: 根据title生成一整篇文章
        """
        pass
    
    @abstractmethod
    def WriteToDocx(self, output_path:str) -> None:
        """
        output_path: 文档输出路径
        作用: 将文章写入word文档
        """
        pass
