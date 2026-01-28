from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableParallel
from operator import itemgetter
from langchain_core.output_parsers import StrOutputParser
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

from docx import *
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import torch
# 数据 理论 实验步骤
#prompt = """请输出一个可以支持“劳动力对经济增长的推动作用”的相关数据表格"""
# llm = ChatOpenAI(model_name="gpt-3.5-turbo",openai_api_key="sk-kFwihnlBAuTgfcZc5d183312767648Ef88C58eFa2bA07b80",temperature=1,base_url = "https://api.oaipro.com/v1") | StrOutputParser()
# table = llm.invoke(prompt)
# print(table)
# torch.save(table, 't')


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

table = torch.load('t')

table = re.sub(' ', '', table)
rows = table.split('\n')

table_list = []
for i in range(len(rows)):
    if i == 1:
        continue
    column_list = rows[i].split('|')
    table_list.append(column_list)
col_num = len(column_list) - 2
row_num = len(table_list)

doc = Document()
table = doc.add_table(rows=row_num, cols=col_num, style='Normal Table')

# 向表格中添加数据
for i in range(row_num):
    for j in range(col_num):
        table.cell(i, j).text = table_list[i][j + 1]
        table.cell(i, j).paragraphs[0].runs[0].font.size = Pt(12)
        table.cell(i, j).paragraphs[0].runs[0].font.name = ''
        table.cell(i, j).paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        if i == 0:
            set_cell_border(table.cell(i, j),
                            top={"sz": 12, "val": "single", "color": "#000000"},
                            bottom={"sz": 12, "val": "single", "color": "#000000"},
                            # left={"sz": 12, "val": "single", "color": "#FFFFFF"},
                            # right={"sz": 12, "val": "single", "color": "#FFFFFF"},
                            )
        if i == row_num - 1:
            set_cell_border(table.cell(i, j),
                bottom={"sz": 12, "val": "single", "color": "#000000"},
                )
 
# 保存Word文档
doc.save('server_lm\paper\\table.docx')
