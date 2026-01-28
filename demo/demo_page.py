from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


# 添加页脚，格式为：第 n 页 共 m 页
def Footer(doc, font_name, font_size):
    # 添加一个节（Section）并获取页脚
    section = doc.sections[0]
    footer = section.footer

    # 在页脚添加一个段落
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置居中对齐

    # 添加文本 "第"
    run1 = paragraph.add_run('第 ')
    font1 = run1.font
    font1.name = font_name
    font1.size = Pt(font_size)   # 设置字体大小
    run1._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # 添加页码相关的字段
    run2 = paragraph.add_run('')
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'begin')
    run2._element.append(fldChar2)

    run3 = paragraph.add_run('')
    fldChar3 = OxmlElement('w:instrText')
    fldChar3.text = 'PAGE'
    font3 = run3.font
    font3.name = 'Times New Roman'
    font3.size = Pt(font_size)   # 设置字体大小
    run3._element.append(fldChar3)

    run4 = paragraph.add_run('')
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'separate')
    run4._element.append(fldChar4)

    run5 = paragraph.add_run('')
    fldChar5 = OxmlElement('w:fldChar')
    fldChar5.set(qn('w:fldCharType'), 'end')
    run5._element.append(fldChar5)

    # 添加文本 "页，共"
    run6 = paragraph.add_run(' 页 共 ')
    font6 = run6.font
    font6.name = font_name
    font6.size = Pt(font_size)   # 设置字体大小
    run6._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    # 添加页数字段
    run7 = paragraph.add_run('')
    fldChar7 = OxmlElement('w:fldChar')
    fldChar7.set(qn('w:fldCharType'), 'begin')
    run7._element.append(fldChar7)

    run8 = paragraph.add_run('')
    fldChar8 = OxmlElement('w:instrText')
    fldChar8.text = 'NUMPAGES'
    font8 = run8.font
    font8.name = 'Times New Roman'
    font8.size = Pt(font_size)   # 设置字体大小
    run8._element.append(fldChar8)

    run9 = paragraph.add_run('')
    fldChar9 = OxmlElement('w:fldChar')
    fldChar9.set(qn('w:fldCharType'), 'separate')
    run9._element.append(fldChar9)

    run10 = paragraph.add_run('')
    fldChar10 = OxmlElement('w:fldChar')
    fldChar10.set(qn('w:fldCharType'), 'end')
    run10._element.append(fldChar10)

    # 添加文本 "页"
    run11 = paragraph.add_run(' 页')
    font11 = run11.font
    font11.name = font_name
    font11.size = Pt(font_size)   # 设置字体大小
    run11._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


# 创建一个新的 Word 文档
doc = Document('test.docx')

Footer(doc=doc, font_name=u'宋体', font_size=10.5)

# 保存文档
doc.save('output_with_footer.docx')

