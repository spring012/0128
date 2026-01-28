import hashlib
import re
 
import docx.document
from docx.oxml import parse_xml
from docx.text.paragraph import Paragraph
 
 
def table_of_contents(doc: docx.document.Document):
    doc.add_heading("目录")
    r = doc.add_paragraph().add_run()
    return r.element
 
 
def add_heading(doc: docx.document.Document, title: str, ml):
    w_id = int(re.findall(r'^(\d+)\.', title)[0])
    bookmark = '_Toc' + hashlib.md5(title.encode('utf-8')).hexdigest()[:4]
    template = f"""<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:pPr><w:pstyle w:val="Heading1"></w:pstyle></w:pPr>
        <w:bookmarkStart w:id="{w_id}" w:name="{bookmark}"/>
        <w:r><w:t>{title}</w:t></w:r>
        <w:bookmarkEnd w:id="{w_id}"/></w:p>
    """
    p = parse_xml(template)
    doc.element.body.insert_element_before(p, 'w:sectPr')
    p = Paragraph(p, doc._body)
    p.style = "Heading 1"
    ml.append(table_of_content(title, 1, bookmark))
 
 
def table_of_content(s_title, page_num, bookmark='_Toc15376'):
    template = f"""<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
        <w:pPr><w:pStyle w:val="26"/><w:tabs><w:tab w:val="right" w:leader="dot" w:pos="8309"/></w:tabs></w:pPr>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> HYPERLINK \\l {bookmark} </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:t>{s_title}</w:t></w:r>
        <w:r><w:tab/></w:r>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText xml:space="preserve"> PAGEREF {bookmark} \\h </w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r><w:t>{page_num}</w:t></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>
    """
    return parse_xml(template)
 
 
doc = docx.Document()
ml = table_of_contents(doc)
 
add_heading(doc, "1.Test1", ml)
add_heading(doc, "2.Test2", ml)
add_heading(doc, "3.Test3", ml)
add_heading(doc, "4.Test4", ml)
add_heading(doc, "5.Test5", ml)
 
doc.save('tes.docx')