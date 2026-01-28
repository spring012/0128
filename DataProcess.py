import os
import re
import fitz
import shutil
import json
import time
from langchain_community.vectorstores.chroma import Chroma
from langchain_milvus.vectorstores import Milvus
from langchain_community.embeddings.huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.document_loaders import DirectoryLoader, PyPDFLoader, PyMuPDFLoader
from Interface import DataBase
from typing import List
import docx 
from langchain_core.documents import Document
import subprocess
import xlrd
import openpyxl
import jieba


def process_file(file_paths: List[str]):
    abstract_regex = re.compile(r'\b摘要\b|\bAbstract\b', flags=re.IGNORECASE)
    content_regex = re.compile(r'\b目录\b|\bContents\b', flags=re.IGNORECASE)
    references_regex = re.compile(r'\b参考文献\b|\bReferences\b', flags=re.IGNORECASE)
    text_splitter = CharacterTextSplitter(chunk_size=1024, chunk_overlap=128, separator='。')

    def process_txt(file_path: str) -> List[str]:
        content = ''
        with open(file_path, 'r', encoding='utf-8') as file:
            for line in file:
                content += line.strip()
        content = re.sub('\n', '', content)
        return text_splitter.split_text(content)

    def process_pdf(file_path: str) -> List[str]:
        with fitz.open(file_path) as doc:
            is_add = 0
            ref_times = 0
            med_page = doc.page_count / 2
            content_list = []
            for index, passage in enumerate(doc.pages()):
                passage_content = passage.get_text()
                is_English = True if passage_content.find('。') == -1 else False

                text = re.sub(' ','', passage_content)
                if abstract_regex.search(text) is not None or content_regex.search(text) is not None:
                    is_add = index
                    content_list = []
                if references_regex.search(text) is not None:
                    ref_times += 1
                    if index < med_page: # 第一次出现在目录里
                        is_add=index
                        content_list = []
                    else: # 最后的参考文献章节
                        break
                
                if is_add < index:
                    content_list.append(passage_content)

        drop_content = ''
        for passage in content_list:
            if is_English:
                passage = passage[passage.find('.') + 1 : passage.rfind('.') + 1]
            else:
                passage = passage[passage.find('。') + 1 : passage.rfind('。') + 1]
                
            passage = re.sub(r'\\u[0-9a-fA-F]{4}', "", passage)  # 去除unicode编码
            passage = re.sub(r'\[\d+\]', '', passage)  # 去掉文献编号
            passage = re.sub(r'[【(（][^)）]*[)）】]', '', passage)  # 去除括号内容
            passage = re.sub(r'\[\d+\]．http[s]?\s*:\s*//\S+\s*|www\.\s*\S+|(?:http[s]?|www)\s*:\s*//\S+\s*', '', passage) # 去除网址
            passage = re.sub(r'[\x00-\x1F\x7F-\xFF]', '', passage)  # 去除控制字符

            for sent in passage.split('.' if is_English else '。'):  # 以句子为单位清洗数据（表格）                       
                if len(re.findall(r'\d+\d+', sent)) < 10 and len(sent) != 0:
                    drop_content +=  sent + '.' if is_English else sent + '。'
        
        return text_splitter.split_text(drop_content)
    
    def process_docx(file_path: str) -> List[str]:
        doc = docx.Document(file_path)
        content = ''
        for paragraph in doc.paragraphs:
            content += paragraph.text
        content = re.sub('\u3000', '', content)  # 去除控制字符
        return text_splitter.split_text(content)

    def process_doc(file_path: str) -> List[str]:
        result = subprocess.run(
            ['soffice', '--headless', '--convert-to', 'docx', file_path, '--outdir', 'tmp'],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True
        )
        if result.returncode == 0:
            file_name = 'tmp\\' + os.path.basename(file_path) + 'x'
            doc = process_docx(file_name)
            os.remove(file_name)
            return doc
        else:
            return []

    def process_xls(file_path: str) -> List[str]:
        xls_data = xlrd.open_workbook(file_path)  
        sheets = []
        for sheetname in xls_data._sheet_names:
            sheet = xls_data.sheet_by_name(sheetname)
            one_sheet = {'Table title': sheetname}
            for row in range(sheet.nrows):  
                if row == 0:
                    one_sheet['Table header'] = [sheet.cell_value(row, col) for col in range(sheet.ncols)]
                else:
                    one_sheet['Row ' + str(row)] = [sheet.cell_value(row, col) for col in range(sheet.ncols)]
            sheets.append(json.dumps(one_sheet, ensure_ascii=False))
        return sheets

    def process_xlsx(file_path: str) -> List[str]:
        workbook = openpyxl.load_workbook(file_path)
        sheets = []
        for sheetname in workbook.sheetnames:
            worksheet = workbook[sheetname]
            one_sheet = {'Table title': sheetname}
            for index, row in enumerate(worksheet.iter_rows()):
                if index == 0:
                    one_sheet['Table header'] = [cell.value for cell in row]
                else:
                    one_sheet['Row ' + str(index)] = [cell.value for cell in row]
            sheets.append(json.dumps(one_sheet, ensure_ascii=False))
        return sheets

    segment = []
    for file_path in file_paths:
        if file_path.endswith('.txt'):
            segment.extend(process_txt(file_path))
        elif file_path.endswith('.pdf'):
            segment.extend(process_pdf(file_path))
        elif file_path.endswith('.docx'):
            segment.extend(process_docx(file_path))
        elif file_path.endswith('.doc'):
            segment.extend(process_doc(file_path))
        elif file_path.endswith('.xlsx'):
            segment.extend(process_xlsx(file_path))
        elif file_path.endswith('.xls'):
            segment.extend(process_xls(file_path))
        else:
            continue
    return segment

class LocalDataBase(DataBase):
    def __init__(self, model_name: str, data_path: str, database_path: str, device: str) -> None:
        """
        model_name:检索器使用模型
        data_path:存放pdf的文件夹路径
        database_path:向量数据库路径
        """
        self.device = device
        self.embedding_model = self.load_model(model_name)
        if os.path.exists(database_path):
            self.embedding_database = Chroma(persist_directory=database_path, embedding_function=self.embedding_model)
        else:
            self.not_use_file = []
            self.embedding_database = self.init_database(data_path, database_path)
            with open('D:\\full-pdfs\\工商管理-最新发表论文倒序-数据库3_11\\not_use\\' + str(time.time()) + '.json', 'w', encoding='utf-8') as f:
                    json.dump({'file': self.not_use_file}, f, ensure_ascii=False, indent=4)

    def load_model(self, model_name):
        embedding_model = HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={"device": self.device},
            encode_kwargs={"normalize_embeddings": True},
        )
        return embedding_model

    def init_database(self, data_path, database_path):
        text_splitter =  CharacterTextSplitter(chunk_size=1024, chunk_overlap=128, separator='。')
        abstract_regex = re.compile(r'\b摘要\b|\bAbstract\b', flags=re.IGNORECASE)
        content_regex = re.compile(r'\b目录\b|\bContents\b', flags=re.IGNORECASE)
        references_regex = re.compile(r'\b参考文献\b|\bReferences\b', flags=re.IGNORECASE)

        documents = []
        for file_index, file_name in enumerate(os.listdir(data_path)):
            with fitz.open(os.path.join(data_path, file_name)) as doc:
                is_add = 0
                ref_times = 0
                med_page = doc.page_count / 2
                content_list = []
                is_English = True
                for index, passage in enumerate(doc.pages()):
                    passage_content = passage.get_text()
                    if passage_content.find('。') != -1: is_English = False

                    text = re.sub(' ','', passage_content)
                    if abstract_regex.search(text) is not None or content_regex.search(text) is not None:
                        is_add = index
                        content_list = []
                    if references_regex.search(text) is not None:
                        ref_times += 1
                        if index < med_page: # 第一次出现在目录里
                            is_add=index
                            content_list = []
                        else: # 最后的参考文献章节
                            break
                    
                    if is_add < index:
                        content_list.append(passage_content)

            drop_content = ''
            for passage in content_list:
                if is_English:
                    passage = passage[passage.find('.') + 1 : passage.rfind('.') + 1]
                else:
                    passage = passage[passage.find('。') + 1 : passage.rfind('。') + 1]
                    
                passage = re.sub(r'\\u[0-9a-fA-F]{4}', "", passage)  # 去除unicode编码
                passage = re.sub(r'\[\d+\]', '', passage)  # 去掉文献编号
                passage = re.sub(r'[【(（][^)）]*[)）】]', '', passage)  # 去除括号内容
                passage = re.sub(r'\[\d+\]．http[s]?\s*:\s*//\S+\s*|www\.\s*\S+|(?:http[s]?|www)\s*:\s*//\S+\s*', '', passage) # 去除网址
                passage = re.sub(r'[\x00-\x1F\x7F-\xFF]', '', passage)  # 去除控制字符

                for sent in passage.split('.' if is_English else '。'):  # 以句子为单位清洗数据（表格）           
                    if len(re.findall(r'\d+\d+', sent)) < 10 and len(sent) != 0:
                        drop_content +=  sent + '.' if is_English else sent + '。'
    
            if len(drop_content) != 0:
                documents.append(Document(drop_content, metadata={'name': file_name.split('.')[0]}))
            else:
                self.not_use_file.append(file_name)
                print(file_index)

        documents_split = text_splitter.split_documents(documents)

        embedding_database = Chroma.from_documents(documents_split, self.embedding_model, persist_directory=database_path)
        embedding_database.persist()
        # embedding_database = Milvus.from_documents(documents_split, self.embedding_model, connection_args={"uri": database_path})
        return embedding_database

    def is_actual_references(self, text):
        # 检查文本中是否包含典型的参考文献格式，例如序号和年份
        pattern = re.compile(r'\[\d+\]\s*(\d{4})')  # 假设参考文献格式为 [序号] 年份
        return bool(pattern.search(text))
        
    def retrieve(self, query: str, tok_k=15):
        docs = self.embedding_database.similarity_search(query, k=tok_k)
        return docs

class ReferenceDataBase(DataBase):
    def __init__(self, model_name: str, pdf_path: str, ref_path: str, ref_database_path: str, device: str) -> None:
        """
        model_name:检索器使用模型
        pdf_path:存放pdf的文件夹路径
        ref_path: 包括引文信息、标题、摘要信息的json文件路径
        database_path:向量数据库路径
        """
        self.device = device
        self.embedding_model = self.load_model(model_name)
        if os.path.exists(ref_database_path):
            self.embedding_database = Chroma(persist_directory=ref_database_path, 
                                             embedding_function=self.embedding_model,
                                             collection_name='all_reference')
        else:
            self.embedding_database = self.init_database(ref_path, ref_database_path)

    def load_model(self, model_name):
        embedding_model = HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={"device": self.device},
            encode_kwargs={"normalize_embeddings": True},
        )
        return embedding_model

    def init_database(self, ref_path, database_path):
        title = []
        # metadata = []
        with open(ref_path, encoding='utf-8') as f:
            title_abstract = json.load(f)

        for item in title_abstract:
            # title.append(item['name'])
            # metadata.append({'abstract':item['abstract'], 'reference':item['reference']})
            title.append(item['abstract'])
        
        embedding_database = Chroma.from_texts(title, self.embedding_model,
                                               persist_directory=database_path,
                                               # metadatas=metadata,
                                               collection_name='all_reference')
        embedding_database.persist()
        return embedding_database

class CreateWordTable():
    def __init__(self, data_path, word_table_path):
        self.data_path = data_path
        self.word_table_path = word_table_path
        
        try:
            with open(word_table_path, encoding='utf-8') as f:
                self.word_table = json.load(f)
        except:
            self.word_table = None
    
    def create_paper_word(self):
        abstract_regex = re.compile(r'\b摘要\b|\bAbstract\b', flags=re.IGNORECASE)
        content_regex = re.compile(r'\b目录\b|\bContents\b', flags=re.IGNORECASE)
        references_regex = re.compile(r'\b参考文献\b|\bReferences\b', flags=re.IGNORECASE)

        word_freq = {}
        for file_index, file_name in enumerate(os.listdir(self.data_path)):
            with fitz.open(os.path.join(self.data_path, file_name)) as doc:
                is_add = 0
                ref_times = 0
                med_page = doc.page_count / 2
                content_list = []
                is_English = True
                for index, passage in enumerate(doc.pages()):
                    passage_content = passage.get_text()
                    if passage_content.find('。') != -1: is_English = False

                    text = re.sub(' ','', passage_content)
                    if abstract_regex.search(text) is not None or content_regex.search(text) is not None:
                        is_add = index
                        content_list = []
                    if references_regex.search(text) is not None:
                        ref_times += 1
                        if index < med_page: # 第一次出现在目录里
                            is_add=index
                            content_list = []
                        else: # 最后的参考文献章节
                            break
                    
                    if is_add < index:
                        content_list.append(passage_content)

            drop_content = ''
            for passage in content_list:
                if is_English:
                    passage = passage[passage.find('.') + 1 : passage.rfind('.') + 1]
                else:
                    passage = passage[passage.find('。') + 1 : passage.rfind('。') + 1]
                    
                passage = re.sub(r'\\u[0-9a-fA-F]{4}', "", passage)  # 去除unicode编码
                passage = re.sub(r'\[\d+\]', '', passage)  # 去掉文献编号
                passage = re.sub(r'[【(（][^)）]*[)）】]', '', passage)  # 去除括号内容
                passage = re.sub(r'\[\d+\]．http[s]?\s*:\s*//\S+\s*|www\.\s*\S+|(?:http[s]?|www)\s*:\s*//\S+\s*', '', passage) # 去除网址
                passage = re.sub(r'[\x00-\x1F\x7F-\xFF]', '', passage)  # 去除控制字符

                for sent in passage.split('.' if is_English else '。'):  # 以句子为单位清洗数据（表格）           
                    if len(re.findall(r'\d+\d+', sent)) < 10 and len(sent) != 0:
                        drop_content +=  sent + '.' if is_English else sent + '。'
    
            if len(drop_content) != 0:
                words = jieba.lcut(drop_content, cut_all=False)
                for word in words:
                    if len(word) <= 1:
                        continue

                    try:
                        word_freq[word] += 1
                    except:
                        word_freq[word] = 1
            print(file_index)
            if file_index % 2000 == 0:
                with open('D:\\full-pdfs\工商管理-最新发表论文倒序-数据库3_11' + str(time.time()) + '.json', 'w', encoding='utf-8') as f:
                    json.dump(word_freq, f, ensure_ascii=False, indent=4)
        
        with open(self.word_table_path, 'w', encoding='utf-8') as f:
            json.dump(word_freq, f, ensure_ascii=False, indent=4)
        self.word_table = word_freq
    
    def clean(self):
        with open('D:\\full-pdfs\工商管理-最新发表论文倒序-数据库3_11\\姓氏.txt', 'r', encoding='utf-8') as file:
            content = file.read().split('\n')
        name = set()
        for item in content:
            if len(item) != 0:
                name.add(item) 
            
        def is_english_or_number(s):
            temp = re.sub('[a-zA-Z]', '', s)
            if len(temp) == 0:
                return True
            else:
                try:
                    float(temp)
                    return True
                except ValueError:
                    return False
        
        def is_percentage(s):
            if '%' in s or '分之' in s:
                return True
            else:
                return False
        
        def is_name(s):
            if s[0] in name:
                return True
            else:
                return False
            
        self.clean_word_table = {}
        for word, freq in self.word_table.items():
            if freq != 1 and not is_english_or_number(word) and not is_percentage(word) and not is_name(word) and len(word) >= 4:
                self.clean_word_table[word] = freq
    
    def create_sys_word(self):
        words = set()
        for file_path in ['D:\\full-pdfs\工商管理-最新发表论文倒序-数据库3_11\\merge_syno.txt', 'D:\\full-pdfs\工商管理-最新发表论文倒序-数据库3_11\\syno_from_baidu_hanyu.txt', 'D:\\full-pdfs\工商管理-最新发表论文倒序-数据库3_11\\syno_from_cwn.txt']:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().split('\n')
            
            for line in content:
                for item in line.split(' '):
                    if len(item) > 1:
                        words.add(item) 
        self.sys_word_tabel = list(words)
        return self.sys_word_tabel

    def bulid_database(self, word_list):
        embedding_model = HuggingFaceEmbeddings(
            model_name= '',
            model_kwargs={"device": 'cuda'},
            encode_kwargs={"normalize_embeddings": True},
        )

        small_word_list = []
        big_word_list = []
        for item in word_list:
            if len(item) >= 4:
                big_word_list.append(item)
            else:
                small_word_list.append(item)
                
        embedding_database = Chroma.from_texts(small_word_list, embedding_model, persist_directory='D:\EmbeddingDatabase\small_word_database')
        embedding_database.persist()

        embedding_database = Chroma.from_texts(big_word_list, embedding_model, persist_directory='D:\EmbeddingDatabase\\big_word_database')
        embedding_database.persist()

if __name__ == '__main__':
    """model_name = 'D:\Model\\text2vec-large-chinese'
    data_path = 'D:\\full-pdfs\工商管理-最新发表论文倒序-数据库3_11\data'
    database_path = 'D:\EmbeddingDatabase'

    A = time.time()
    EmbeddingBase = LocalDataBase(model_name, data_path, database_path, device='cuda')
    B = time.time()
    result = EmbeddingBase.retrieve(query='劳动赋予人的意义')
    C = time.time()
    print(result)
    print(B - A, C - B)"""

    data_path = 'D:\\full-pdfs\工商管理-最新发表论文倒序-数据库3_11\data'
    word_table_path = 'D:\\full-pdfs\工商管理-最新发表论文倒序-数据库3_11\\word_freq.json'
    word_table = CreateWordTable(data_path, word_table_path)
    # word_table.create_paper_word()
    # word_table.clean()
    table = word_table.create_sys_word()
    word_table.bulid_database(table)
