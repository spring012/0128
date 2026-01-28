import random
import numpy as np
import torch
import json
from collections import defaultdict
from model import load_tokenizer, load_model
from fast_detect_gpt import get_sampling_discrepancy_analytic
import glob
import os
import argparse
import re
from docx import Document
from docx.shared import RGBColor

# 估计AIGC生成概率
class ProbEstimator:
    def __init__(self, args):
        self.real_crits = []
        self.fake_crits = []
        for result_file in glob.glob(os.path.join(args.ref_path, '*.json')):
            with open(result_file, 'r') as fin:
                res = json.load(fin)
                self.real_crits.extend(res['predictions']['real'])
                self.fake_crits.extend(res['predictions']['samples'])
        print(f'ProbEstimator: total {len(self.real_crits) + len(self.fake_crits)} samples.')

    def crit_to_prob(self, crit):
        all_crits = np.array(self.real_crits + self.fake_crits)
        if all_crits.size == 0:
            raise ValueError("real_crits 和 fake_crits 均为空，无法计算概率。")

        sorted_offsets = np.sort(np.abs(all_crits - crit))
        if sorted_offsets.size <= 100:
            raise IndexError("数组大小不足，无法访问索引 100。")

        offset = sorted_offsets[100]
        cnt_real = np.sum((np.array(self.real_crits) > crit - offset) & (np.array(self.real_crits) < crit + offset))
        cnt_fake = np.sum((np.array(self.fake_crits) > crit - offset) & (np.array(self.fake_crits) < crit + offset))

        return cnt_fake / (cnt_real + cnt_fake) if (cnt_real + cnt_fake) > 0 else 0.0

def split_text_by_sentences(text):
    sentences = re.split(r'(。|！|？|；)', text)  # 以标点符号分割
    result = [sentences[i] + sentences[i+1] if i+1 < len(sentences) else sentences[i]
              for i in range(0, len(sentences), 2)]
    return [s.strip() for s in result if s.strip()]


# 通过滑动窗口分割文本，每次重叠 2 个句子
def split_text_by_sentences_with_overlap(sentences, block_size=8, overlap=2):
    start = 0
    while start + block_size <= len(sentences):
        yield sentences[start:start + block_size]
        start += (block_size - overlap)  # 每次滑动时重叠2个句子
    if start < len(sentences):  # 确保最后的句子也被处理
        yield sentences[start:]

def is_title(paragraph):
    return bool(re.match(r"^\d+(\.\d+)*\s+", paragraph)) or len(paragraph) < 20

def get_aigc_interval(prob):
    if prob < 0.6:
        return "其他"
    elif prob < 0.8:
        return "轻度疑似"
    elif prob < 0.9:
        return "中度疑似"
    else:
        return "高度疑似"

def compute_aigc_prob(block, scoring_tokenizer, scoring_model, reference_tokenizer, reference_model, criterion_fn, prob_estimator, device):
    block_text = "".join(block)
    tokenized = scoring_tokenizer(block_text, truncation=True, return_tensors="pt", padding=True, return_token_type_ids=False).to(device)
    labels = tokenized.input_ids[:, 1:]

    with torch.no_grad():
        logits_score = scoring_model(**tokenized).logits[:, :-1]
        if reference_model is None:
            logits_ref = logits_score
        else:
            tokenized_ref = reference_tokenizer(block_text, truncation=True, return_tensors="pt", padding=True, return_token_type_ids=False).to(device)
            logits_ref = reference_model(**tokenized_ref).logits[:, :-1]

        crit = criterion_fn(logits_ref, logits_score, labels)

    prob = prob_estimator.crit_to_prob(crit)
    del tokenized, logits_score, logits_ref
    torch.cuda.empty_cache()

    return prob

def detect_aigc_in_text(text, args):
    scoring_tokenizer = load_tokenizer(args.scoring_model_name, args.dataset, args.cache_dir)
    scoring_model = load_model(args.scoring_model_name, args.device, args.cache_dir)
    scoring_model.eval()

    reference_tokenizer, reference_model = None, None
    if args.reference_model_name != args.scoring_model_name:
        reference_tokenizer = load_tokenizer(args.reference_model_name, args.dataset, args.cache_dir)
        reference_model = load_model(args.reference_model_name, args.device, args.cache_dir)
        reference_model.eval()

    criterion_fn = get_sampling_discrepancy_analytic
    prob_estimator = ProbEstimator(args)

    sentences = split_text_by_sentences(text)
    sentence_probs = defaultdict(list)

    # 使用滑动窗口处理文本
    for block in split_text_by_sentences_with_overlap(sentences, block_size=8, overlap=2):
        prob = compute_aigc_prob(block, scoring_tokenizer, scoring_model, reference_tokenizer, reference_model, criterion_fn, prob_estimator, args.device)
        # 对窗口中的每个句子，给出一个平均概率
        for sentence in block:
            sentence_probs[sentence].append(prob)

    # 最终的句子概率
    final_sentence_probs = [(sentence, np.mean(probs)) for sentence, probs in sentence_probs.items()]
    return merge_and_annotate_results(final_sentence_probs)

# 在docx中标注结果
def merge_and_annotate_results(sentence_probs):
    merged_results = []
    current_block = []
    current_interval = None

    for sentence, prob in sentence_probs:
        interval = get_aigc_interval(prob)

        # 合并具有相同区间的句子
        if current_interval is None or interval == current_interval:
            current_block.append((sentence, prob))
            current_interval = interval
        else:
            merged_results.append((current_block, current_interval))
            current_block = [(sentence, prob)]
            current_interval = interval

    # 最后一部分
    if current_block:
        merged_results.append((current_block, current_interval))

    annotated_text = ""
    for block, interval in merged_results:
        block_text = " ".join(sentence for sentence, _ in block)
        avg_prob = np.mean([prob for _, prob in block])
        if interval!='其他':
            annotated_text += f"{block_text} [{interval},{avg_prob * 100:.2f}%AIGC]|"
        else:
            annotated_text += f"{block_text}|"

    return annotated_text

def detect_aigc_in_file(file_path, args):
    if os.path.exists(file_path):
        # 读取 DOCX 文件内容
        document = Document(file_path)
        # 获取非空段落并将它们连接为一个文本
        text = "\n".join([para.text for para in document.paragraphs if para.text.strip()])
    else:
        raise FileNotFoundError(f"文件 {file_path} 不存在！")
    return text,detect_aigc_in_text(text, args)

def save_annotated_result(original_text,result_text, output_path):
    # 检查路径是否存在，如果不存在则创建
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 创建 DOCX 文档
    doc = Document()
    paragraphs = original_text.split("\n")  # 论文原始段落
    results = result_text.split("\n")  # 检测后的段落

    # for para_text, result in zip(paragraphs, results):
    #     if is_title(para_text):  # 跳过标题部分，不检测
    #         doc.add_paragraph(para_text)
    #         continue

    # 逐段添加带标注的文本
    for paragraph in result_text.split("|"):
        para = doc.add_paragraph()
        run = para.add_run(paragraph)  # 将段落文本添加到一个run中
        if "轻度疑似" in paragraph:
            run.font.color.rgb = RGBColor(255, 225, 0)  
        elif "中度疑似" in paragraph:
            run.font.color.rgb = RGBColor(255, 205, 0)  
        elif "高度疑似" in paragraph:
            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        run.font.name = '微软雅黑'  
    # 保存文档
    doc.save(output_path)
    print(f"检测结果已保存到 {output_path}")

